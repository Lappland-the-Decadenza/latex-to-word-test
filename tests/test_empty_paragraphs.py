import docx

from docfidelity import compare
from latexword.docx.read import docx_to_latex
from latexword.docx.write import convert_latex_to_docx


def test_empty_paragraph_survives_round_trip_and_is_measured(tmp_path):
    source = tmp_path / "source.docx"
    original = docx.Document()
    original.add_paragraph("Before")
    original.add_paragraph()
    original.add_paragraph("After")
    original.save(source)

    tex_path = tmp_path / "roundtrip.tex"
    tex, warnings = docx_to_latex(str(source), str(tex_path))
    assert not warnings
    assert r"\mbox{}\par" in tex
    tex_path.write_text(tex, encoding="utf-8")

    output = tmp_path / "output.docx"
    _, warnings = convert_latex_to_docx(str(tex_path), str(output))
    assert not warnings
    assert [p.text for p in docx.Document(output).paragraphs] == [
        "Before", "", "After",
    ]

    without_blank = tmp_path / "without-blank.docx"
    shortened = docx.Document()
    shortened.add_paragraph("Before")
    shortened.add_paragraph("After")
    shortened.save(without_blank)
    findings = compare(str(source), str(without_blank))
    assert any(
        finding.verdict == "degradation"
        and "block lost: ''" in finding.detail
        for finding in findings
    )
