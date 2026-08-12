"""PLAN.md §7.1 -- footnotes and endnotes, both directions.

Before this change, ``word/footnotes.xml`` was never read and ``\\footnote``
was an unknown command: the note's text vanished entirely on a round trip
(measured: a corpus document's "1 footnotes 1 -> 0" degradation), and the
plan names notes the single highest-value item -- they are text, not
decoration.

These tests exercise the forward and reverse halves independently plus one
true round trip through the real pipeline (the same shape as the smoke
that drove the implementation), rather than depending on the private
corpus documents the fidelity report measures against.
"""

import os
import sys
import tempfile
import zipfile

import docx
from lxml import etree

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from latexword.docx import read as docx_read
from latexword.docx import write as latex2word

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"


def _real_notes(path, part):
    """Real note definitions of one part: id -> text (see docfidelity)."""
    with zipfile.ZipFile(path) as z:
        if part not in z.namelist():
            return {}
        root = etree.fromstring(z.read(part))
    tag = "footnote" if "footnotes" in part else "endnote"
    out = {}
    for el in root.iter(W + tag):
        if el.get(W + "type") is not None:
            continue
        nid = el.get(W + "id")
        if nid is not None:
            out[nid] = "".join(t.text or "" for t in el.iter(W + "t"))
    return out


def _build(text, warnings=None):
    """The real pipeline: .tex file -> .docx.

    Accepts a body fragment or a complete document (the reverse's output is
    a complete document; wrapping it again would leak its preamble into the
    body of the next generation)."""
    if "\\documentclass" not in text:
        text = ("\\documentclass{article}\n"
                "\\begin{document}\n" + text + "\n"
                "\\end{document}\n")
    tmp = tempfile.mkdtemp()
    tp = os.path.join(tmp, "src.tex")
    with open(tp, "w", encoding="utf-8") as fh:
        fh.write(text)
    out = os.path.join(tmp, "g1.docx")
    latex2word.convert_latex_to_docx(tp, out)
    return out


# --- Forward: \footnote{...} creates a real notes part ----------------------


def test_footnote_wires_part_content_type_and_relationship():
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), r"a\footnote{note} b")
    path = os.path.join(tempfile.mkdtemp(), "fn.docx")
    d.save(path)
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        ct = z.read("[Content_Types].xml").decode("utf-8")
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        docroot = etree.fromstring(z.read("word/document.xml"))
    assert "word/footnotes.xml" in names
    assert "footnotes+xml" in ct
    assert 'Target="footnotes.xml"' in rels
    refs = list(docroot.iter(W + "footnoteReference"))
    assert len(refs) == 1
    # the mark must be a superscripted run: Word renders the reference with
    # the character style normally, but a missing style must not leave a
    # plain-size digit.
    va = refs[0].getparent().find(W + "rPr/" + W + "vertAlign")
    assert va is not None and va.get(W + "val") == "superscript"


def test_footnote_body_survives_with_formatting_and_math():
    d = docx.Document()
    latex2word.add_inline_latex(
        d.add_paragraph(), r"a\footnote{see \textbf{note} and $x^2$} b")
    path = os.path.join(tempfile.mkdtemp(), "fn.docx")
    d.save(path)
    notes = _real_notes(path, "word/footnotes.xml")
    assert len(notes) == 1
    text = next(iter(notes.values()))
    assert "see note and" in text  # \textbf is formatting; text survives
    # the oMath object is inside the notes part, not dropped
    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read("word/footnotes.xml"))
    assert root.find(".//" + "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath") is not None


def test_hyperlink_in_note_body_keeps_its_namespace():
    # Defect that was silent: the notes tree's root declared no namespaces,
    # lxml invented prefixes at serialization, and the r:id attribute
    # collided with the generated prefix -- the serializer rebound the
    # prefix on the hyperlink element itself, changing its tag to the
    # relationships namespace. The w:t under it was then read as
    # {relationships}t and the link text vanished from the note.
    d = docx.Document()
    latex2word.add_inline_latex(
        d.add_paragraph(), r"a\footnote{see \href{https://example.com}{link}} b")
    path = os.path.join(tempfile.mkdtemp(), "fn.docx")
    d.save(path)
    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read("word/footnotes.xml"))
    hlinks = root.iter(W + "hyperlink")
    assert sum(1 for _ in hlinks) == 1
    notes = _real_notes(path, "word/footnotes.xml")
    assert "link" in next(iter(notes.values()))


def test_nested_footnote_is_plain_text_with_a_warning():
    # LaTeX forbids nested notes; writing a reference inside footnotes.xml
    # is invalid OOXML Word would reject.
    d = docx.Document()
    w = []
    latex2word.add_inline_latex(d.add_paragraph(), r"a\footnote{outer \footnote{inner}}", warnings=w)
    path = os.path.join(tempfile.mkdtemp(), "fn.docx")
    d.save(path)
    assert any("nested" in x for x in w)
    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read("word/footnotes.xml"))
    assert list(root.iter(W + "footnoteReference")) == []


def test_endnote_writes_endnotes_part_and_marker_is_consumed():
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), r"a\endnote{an endnote} b")
    path = os.path.join(tempfile.mkdtemp(), "en.docx")
    d.save(path)
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
    assert "word/endnotes.xml" in names
    assert 'Target="endnotes.xml"' in rels
    assert len(_real_notes(path, "word/endnotes.xml")) == 1
    # \theendnotes (which our reverse writes before \end{document}) must be
    # consumed silently, not reported as a dropped unknown command.
    d = docx.Document()
    w = []
    latex2word.add_inline_latex(d.add_paragraph(), "\\theendnotes", warnings=w)
    assert w == []


# --- Reverse: the notes part becomes \footnote{...} -------------------------


def test_reverse_emits_footnote_with_full_body():
    d = docx.Document()
    latex2word.add_inline_latex(
        d.add_paragraph(), r"before \footnote{see \textbf{note} and $x^{2}$ and \href{https://example.com}{link}} after")
    path = os.path.join(tempfile.mkdtemp(), "fn.docx")
    d.save(path)
    tex, warnings = docx_read.docx_to_latex(path)
    assert "\\footnote{see \\textbf{note} and $x^{2}$ and \\href{https://example.com}{link}}" in tex
    assert warnings == []


def test_reverse_dangling_reference_is_named():
    # A reference whose definition is missing (source document inconsistent)
    # must warn, not silently vanish or crash.
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), r"t\footnote{real}")
    path = os.path.join(tempfile.mkdtemp(), "fn.docx")
    d.save(path)
    # break the document: point the reference at a nonexistent id
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(path + ".b", "w") as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name == "word/document.xml":
                root = etree.fromstring(data)
                ref = root.find(".//" + W + "footnoteReference")
                ref.set(W + "id", "99")
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(name, data)
    tex, warnings = docx_read.docx_to_latex(path + ".b")
    assert any("no definition" in w for w in warnings)


def test_reverse_multiparagraph_note_is_named():
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), r"a\footnote{one two} b")
    path = os.path.join(tempfile.mkdtemp(), "fn.docx")
    d.save(path)
    # split the note's paragraph into two paragraphs (move the last run)
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(path + ".b", "w") as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name == "word/footnotes.xml":
                root = etree.fromstring(data)
                note = next(el for el in root.iter(W + "footnote")
                            if el.get(W + "type") is None)
                p = note.find(W + "p")
                runs = p.findall(W + "r")
                p2 = etree.Element(W + "p")
                p2.append(runs[-1])  # moves the run: the note now spans two paragraphs
                p.addnext(p2)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(name, data)
    tex, warnings = docx_read.docx_to_latex(path + ".b")
    assert "\\footnote{" in tex and "one" in tex and "two" in tex
    assert any("paragraph" in w for w in warnings)


# --- Round trip through the real pipeline -----------------------------------


def test_footnote_round_trip_is_stable():
    src = r"Body\footnote{n1} more \footnote{n2}."
    out1 = _build(src)
    tex1, w1 = docx_read.docx_to_latex(out1)
    assert "\\footnote{n1}" in tex1 and "\\footnote{n2}" in tex1
    assert w1 == []
    out2 = _build(tex1)
    tex2, w2 = docx_read.docx_to_latex(out2)
    assert tex2 == tex1  # fixed point from generation 2
    assert w2 == []


def test_endnote_round_trip_emits_preamble_pieces():
    out1 = _build(r"Body\endnote{an endnote}.")
    tex1, w1 = docx_read.docx_to_latex(out1)
    assert "\\endnote{an endnote}" in tex1
    assert "\\usepackage{endnotes}" in tex1
    assert "\\theendnotes" in tex1
    assert w1 == []
    # the reverse output must convert back without the preamble leaking
    # into the body and without note-related warnings
    out2 = _build(tex1)
    tex2, w2 = docx_read.docx_to_latex(out2)
    assert tex2 == tex1
    assert w2 == []
