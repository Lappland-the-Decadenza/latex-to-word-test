"""PLAN.md §7.1 -- cross-references and bookmarks, both directions.

Before this change ``w:bookmarkStart`` fell through the inline loop and the
fidelity report measured 107 "bookmarks N -> 0" degradations (90 of them
Word- or converter-generated names), and ``\\ref{...}`` degraded to
bracketed placeholder text. Word fields were read run-by-run: a REF
field's cached number came back as stale text and a HYPERLINK field lost
its URL entirely (the instruction runs carry no ``w:t``, so nothing in the
old loop even noticed them).

These tests pin the grammar rules: which bookmark families are Word
bookkeeping (consumed silently), which field instructions resolve to what
(and what happens to the cached result), plus one true round trip.
"""

import os
import sys
import tempfile

import docx
from docx import Document
from lxml import etree

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from latexword.docx import read as docx_read
from latexword.docx import write as latex2word

_W_URI = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = "{" + _W_URI + "}"


def _build(text):
    """The real pipeline: .tex file -> .docx. Accepts a body fragment or a
    complete document."""
    if "\\documentclass" not in text:
        text = ("\\documentclass{article}\n"
                "\\begin{document}\n" + text + "\n"
                "\\end{document}\n")
    tmp = tempfile.mkdtemp()
    tp = os.path.join(tmp, "src.tex")
    with open(tp, "w", encoding="utf-8") as fh:
        fh.write(text)
    out = os.path.join(tmp, "g1.docx")
    latex2word.convert_latex_to_docx(tp, out)
    return out


def _docx_with(text, warnings=None):
    """A one-paragraph docx built through the real forward path."""
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), text, warnings=warnings)
    path = os.path.join(tempfile.mkdtemp(), "x.docx")
    d.save(path)
    return path


def _field_runs(instr, cached):
    """The runs of one hand-written field (begin/instr/separate/cached/end)
    -- the shape Word writes -- ready to append to a paragraph element."""
    p = etree.fromstring(
        f'<w:p xmlns:w="{_W_URI}">'
        f"<w:r><w:fldChar w:fldCharType=\"begin\"/></w:r>"
        f"<w:r><w:instrText xml:space=\"preserve\">{instr}</w:instrText></w:r>"
        f"<w:r><w:fldChar w:fldCharType=\"separate\"/></w:r>"
        f"<w:r><w:t>{cached}</w:t></w:r>"
        f"<w:r><w:fldChar w:fldCharType=\"end\"/></w:r>"
        f"</w:p>")
    return list(p)


def _reverse_of(path):
    """docx -> (tex, warnings)."""
    return docx_read.docx_to_latex(path)


def _bookmark_names(path):
    doc = Document(path)
    return [el.get(W + "name") for el in doc.element.iter(W + "bookmarkStart")]


def _body(tex):
    """The tex without its preamble, so assertions cannot match the
    \newcommand declarations (which legitimately contain "1"/"2")."""
    return tex[tex.index("\\begin{document}") + len("\\begin{document}"):]


def _instr_text(path):
    doc = Document(path)
    return "".join(t.text or "" for t in doc.element.iter(W + "instrText"))


def _fld_types(path):
    doc = Document(path)
    return [fc.get(W + "fldCharType") for fc in doc.element.iter(W + "fldChar")]


# --- Forward: \label / \ref / \pageref / \eqref ----------------------------


def test_forward_label_writes_zero_length_bookmark():
    path = _docx_with(r"see \label{fig:one} text")
    doc = Document(path)
    starts = list(doc.element.iter(W + "bookmarkStart"))
    ends = list(doc.element.iter(W + "bookmarkEnd"))
    assert len(starts) == 1 and len(ends) == 1
    assert starts[0].get(W + "name") == "fig:one"
    # zero-length: start and end adjacent, same id, nothing between
    p = starts[0].getparent()
    si = list(p).index(starts[0])
    assert list(p)[si + 1] is ends[0]
    assert starts[0].get(W + "id") == ends[0].get(W + "id")


def test_forward_label_unrepresentable_name_is_named():
    w = []
    path = _docx_with(r"a \label{two words} b", warnings=w)
    assert any("cannot be referenced" in x for x in w)
    assert _bookmark_names(path) == []


def test_forward_ref_writes_field_with_question_cached_result():
    path = _docx_with(r"see \ref{eq:1}.")
    assert _instr_text(path).strip().startswith("REF eq:1")
    assert _fld_types(path) == ["begin", "separate", "end"]
    # The cached result is "?" -- the LaTeX display for an unresolved
    # reference: Word cannot know the number, and a fabricated one would
    # be wrong content. It sits in the separate fldChar's own run.
    doc = Document(path)
    separate = next(fc for fc in doc.element.iter(W + "fldChar")
                    if fc.get(W + "fldCharType") == "separate")
    t = separate.getparent().find(W + "t")
    assert t is not None and t.text == "?"


def test_forward_pageref_and_eqref():
    p1 = _docx_with(r"\pageref{sec:2}")
    p2 = _docx_with(r"\eqref{eq:3}")
    for path, kind, name in ((p1, "PAGEREF", "sec:2"), (p2, "REF", "eq:3")):
        assert _instr_text(path).strip().startswith(f"{kind} {name}")


def test_forward_cite_stays_bracket_placeholder():
    path = _docx_with(r"as \cite{smith2020} shows")
    assert _fld_types(path) == []
    doc = Document(path)
    text = "".join(t.text or "" for t in doc.element.iter(W + "t"))
    assert "[smith2020]" in text


# --- Reverse: bookmarks ----------------------------------------------------


def test_reverse_authored_bookmark_emits_label_at_anchor():
    path = _docx_with(r"before \label{fig:one} after")
    tex, warnings = _reverse_of(path)
    assert "\\label{fig:one}" in tex
    assert warnings == []
    # the anchor sits between the surrounding words, in document order
    assert tex.index("before") < tex.index("\\label{fig:one}") < tex.index("after")


def test_reverse_word_internal_bookmark_families_are_silent():
    # _GoBack, _TocNNN, _HlkNNN and converter X-hex names are Word's own
    # bookkeeping (same class as rsid/proofErr): no \label, no warning.
    path = _docx_with(r"a\label{_GoBack}\label{_Toc123456789}\label{_Hlk987654321}"
                      r"\label{Xc67fd51f63966fd4c64b3427ee905109ad9d407} b")
    tex, warnings = _reverse_of(path)
    assert "\\label{" not in tex
    assert warnings == []


# --- Reverse: fields -------------------------------------------------------


def test_reverse_ref_field_emits_ref_and_drops_cached_number():
    # The cached number is stale (Word froze it at the last update); the
    # \ref regenerates it, so emitting it would double the number.
    d = docx.Document()
    p = d.add_paragraph("see ")
    p._element.extend(_field_runs(" REF _Ref360318195 \\h \\* MERGEFORMAT ", "1.5"))
    path = os.path.join(tempfile.mkdtemp(), "x.docx")
    d.save(path)
    tex, warnings = _reverse_of(path)
    assert "\\ref{_Ref360318195}" in tex
    assert "1.5" not in _body(tex)
    assert warnings == []


def test_reverse_pageref_field():
    d = docx.Document()
    p = d.add_paragraph("p. ")
    p._element.extend(_field_runs(" PAGEREF _Ref123 \\h \\* MERGEFORMAT ", "4"))
    path = os.path.join(tempfile.mkdtemp(), "x.docx")
    d.save(path)
    tex, warnings = _reverse_of(path)
    assert "\\pageref{_Ref123}" in tex
    assert "4" not in _body(tex)


def test_reverse_hyperlink_field_uses_native_href():
    # The field form of a link: instruction carries the URL, cached text is
    # the display. Before this change the URL vanished entirely.
    d = docx.Document()
    p = d.add_paragraph("see ")
    p._element.extend(_field_runs(' HYPERLINK "https://example.com" ', "the site"))
    path = os.path.join(tempfile.mkdtemp(), "x.docx")
    d.save(path)
    tex, warnings = _reverse_of(path)
    assert r"\href{https://example.com}{the site}" in tex
    assert _body(tex).count("the site") == 1
    assert warnings == []


def test_reverse_split_instruction_hyperlink_field():
    # Measured in the corpus: Word splits a long field instruction across
    # several instrText runs, one token per run, spaces stripped at the
    # boundaries. The instruction is their concatenation; before this
    # test the URL came back with a space at every fragment boundary
    # (the fragments were joined with " ") and the link was broken.
    d = docx.Document()
    p = d.add_paragraph("see ")
    frags = ["HYPERLINK", '"', "https", "://", "example", ".", "com",
             "/a", "b", '"']
    runs = [etree.fromstring(
        f'<w:r xmlns:w="{_W_URI}"><w:instrText>{f}</w:instrText></w:r>')
        for f in frags]
    p._element.extend([etree.fromstring(
        f'<w:r xmlns:w="{_W_URI}"><w:fldChar w:fldCharType="begin"/></w:r>')]
        + runs + [etree.fromstring(
        f'<w:r xmlns:w="{_W_URI}"><w:fldChar w:fldCharType="separate"/></w:r>'),
        etree.fromstring(
        f'<w:r xmlns:w="{_W_URI}"><w:t>the site</w:t></w:r>'),
        etree.fromstring(
        f'<w:r xmlns:w="{_W_URI}"><w:fldChar w:fldCharType="end"/></w:r>')])
    path = os.path.join(tempfile.mkdtemp(), "x.docx")
    d.save(path)
    tex, warnings = _reverse_of(path)
    assert r"\href{https://example.com/ab}{the site}" in tex
    assert "the site" in tex
    assert warnings == []


def test_reverse_seq_field_keeps_display_and_reports_loss():
    # Figure-numbering fields have no clean L* spelling, so the complete
    # instruction and cached display use the grammar-safe carrier.
    d = docx.Document()
    p = d.add_paragraph()
    p._element.extend(_field_runs(" SEQ Figure \\* ARABIC ", "1"))
    path = os.path.join(tempfile.mkdtemp(), "x.docx")
    d.save(path)
    tex, warnings = _reverse_of(path)
    assert "1" in _body(tex)
    assert "\\begin{" not in _body(tex)
    assert any("SEQ" in warning for warning in warnings)


def test_reverse_unknown_field_keeps_display_and_reports_loss():
    d = docx.Document()
    p = d.add_paragraph()
    p._element.extend(_field_runs(" NUMPAGES ", "12"))
    path = os.path.join(tempfile.mkdtemp(), "x.docx")
    d.save(path)
    tex, warnings = _reverse_of(path)
    assert "12" in _body(tex)
    assert "\\begin{" not in _body(tex)
    assert any("field instruction" in warning for warning in warnings)


def test_reverse_fldsimple_resolves_like_run_group():
    d = docx.Document()
    p = d.add_paragraph("a ")
    fld = etree.fromstring(
        f'<w:fldSimple xmlns:w="{_W_URI}" w:instr=" REF tab:1 ">'
        f"<w:r><w:t>2</w:t></w:r></w:fldSimple>")
    p._element.append(fld)
    path = os.path.join(tempfile.mkdtemp(), "x.docx")
    d.save(path)
    tex, warnings = _reverse_of(path)
    assert "\\ref{tab:1}" in tex and "2" not in _body(tex)


# --- Round trip through the real pipeline ----------------------------------


def test_bookmark_and_ref_round_trip_is_stable():
    src = r"Section \label{sec:intro} with \ref{sec:intro} and \eqref{eq:a}."
    out1 = _build(src)
    tex1, w1 = _reverse_of(out1)
    assert "\\label{sec:intro}" in tex1
    assert "\\ref{sec:intro}" in tex1
    assert "\\eqref" not in tex1  # canonicalised: every REF field is \ref
    assert w1 == []
    out2 = _build(tex1)
    tex2, w2 = _reverse_of(out2)
    assert tex2 == tex1  # fixed point from generation 2
    assert w2 == []


def test_math_label_still_tolerated_with_warning():
    # A bookmark is a paragraph-level element; it cannot sit inside an
    # OMML equation, so a label in math keeps its named drop (§7.1) -- and
    # the equation itself survives.
    d = docx.Document()
    w = []
    latex2word.add_inline_latex(d.add_paragraph(), r"$x^2 \label{eq:1}$", warnings=w)
    path = os.path.join(tempfile.mkdtemp(), "x.docx")
    d.save(path)
    assert any("label" in x for x in w)
    assert _bookmark_names(path) == []
    m = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
    doc = Document(path)
    assert len(list(doc.element.iter(m + "oMath"))) == 1
