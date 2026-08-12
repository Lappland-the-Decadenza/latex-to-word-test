"""Phase A package integrity and reference numbering regression."""

import os
import zipfile
import xml.etree.ElementTree as std_etree

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from latexword.docx.package import (
    STYLES_WITH_EFFECTS_RELTYPE,
    validate_docx_package,
)
from latexword.docx.read import docx_to_latex
from latexword.docx.write import convert_latex_to_docx


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _numbering_definition(document, abstract_id, num_id):
    root = document.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    level.append(text)
    abstract.append(level)
    first_num = root.find(qn("w:num"))
    if first_num is None:
        root.append(abstract)
    else:
        first_num.addprevious(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    number.append(abstract_ref)
    root.append(number)


def _style_numbering(style, num_id):
    ppr = style._element.get_or_add_pPr()
    numpr = ppr.get_or_add_numPr()
    numpr.get_or_add_numId().val = num_id


def test_copy_mode_preserves_reference_numbering_and_allocates_after_it(tmp_path):
    reference = docx.Document()
    first = reference.styles.add_style("ReferenceNumberedA", WD_STYLE_TYPE.PARAGRAPH)
    second = reference.styles.add_style("ReferenceNumberedB", WD_STYLE_TYPE.PARAGRAPH)
    _style_numbering(first, 11)
    _style_numbering(second, 27)
    _numbering_definition(reference, 11, 11)
    _numbering_definition(reference, 27, 27)
    for rel in list(reference.part.rels.values()):
        if rel.reltype == STYLES_WITH_EFFECTS_RELTYPE:
            del reference.part.rels[rel.rId]
    reference_path = os.path.join(str(tmp_path), "reference.docx")
    reference.save(reference_path)

    tex_path = os.path.join(str(tmp_path), "input.tex")
    with open(tex_path, "w", encoding="utf-8") as stream:
        stream.write(
            r"first" "\n"
            r"second" "\n"
            r"\begin{itemize}\item outer\begin{enumerate}\item inner\end{enumerate}\end{itemize}"
        )
    output_path = os.path.join(str(tmp_path), "output.docx")

    convert_latex_to_docx(
        tex_path,
        output_path,
        reference_path,
        reference_mode="copy",
    )

    assert validate_docx_package(output_path) == []
    with zipfile.ZipFile(output_path) as archive:
        assert "word/stylesWithEffects.xml" not in archive.namelist()
        numbering = std_etree.fromstring(archive.read("word/numbering.xml"))
        styles = std_etree.fromstring(archive.read("word/styles.xml"))

    num_ids = {
        int(element.get(W + "numId"))
        for element in numbering.findall(W + "num")
    }
    assert {11, 27}.issubset(num_ids)
    assert {"ReferenceNumberedA", "ReferenceNumberedB"}.issubset({
        element.get(W + "styleId") for element in styles.findall(W + "style")
    })

    output = docx.Document(output_path)
    generated_list_ids = {
        int(paragraph._p.pPr.numPr.numId.val)
        for paragraph in output.paragraphs
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
    }
    assert generated_list_ids
    assert all(num_id > 27 for num_id in generated_list_ids)


def test_template_numbered_style_is_not_overridden_by_direct_list_numbering(tmp_path):
    reference = docx.Document()
    style = reference.styles.add_style("TemplateNumbered", WD_STYLE_TYPE.PARAGRAPH)
    _style_numbering(style, 11)
    _numbering_definition(reference, 11, 11)
    reference.add_paragraph("styled item")._p.style = "TemplateNumbered"
    reference_path = os.path.join(str(tmp_path), "reference.docx")
    reference.save(reference_path)

    tex_path = os.path.join(str(tmp_path), "input.tex")
    clean_tex, reverse_warnings = docx_to_latex(
        reference_path, tex_path=tex_path
    )
    assert reverse_warnings == []
    with open(tex_path, "w", encoding="utf-8") as stream:
        stream.write(clean_tex)
    output_path = os.path.join(str(tmp_path), "output.docx")

    convert_latex_to_docx(
        tex_path, output_path, reference_path, reference_mode="copy"
    )

    paragraph = docx.Document(output_path).paragraphs[0]
    assert paragraph.style.style_id == "TemplateNumbered"
    assert paragraph._p.pPr.numPr is None


def test_synthetic_bullet_levels_use_glyphs_present_in_their_fonts(tmp_path):
    tex_path = os.path.join(str(tmp_path), "input.tex")
    with open(tex_path, "w", encoding="utf-8") as stream:
        stream.write(
            r"\begin{itemize}\item outer"
            r"\begin{itemize}\item inner\end{itemize}"
            r"\end{itemize}"
        )
    output_path = os.path.join(str(tmp_path), "output.docx")

    convert_latex_to_docx(tex_path, output_path)

    with zipfile.ZipFile(output_path) as archive:
        numbering = std_etree.fromstring(archive.read("word/numbering.xml"))
    bullet_levels = {}
    for level in numbering.iter(W + "lvl"):
        num_fmt = level.find(W + "numFmt")
        if num_fmt is None or num_fmt.get(W + "val") != "bullet":
            continue
        text = level.find(W + "lvlText").get(W + "val")
        fonts = level.find(W + "rPr/" + W + "rFonts")
        bullet_levels[level.get(W + "ilvl")] = (text, fonts.get(W + "ascii"))

    assert bullet_levels["0"] == ("\uf0b7", "Symbol")
    assert bullet_levels["1"] == ("o", "Courier New")
    assert bullet_levels["2"] == ("\uf0a7", "Wingdings")
