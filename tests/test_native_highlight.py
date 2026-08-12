import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

from latexword.docx.read import docx_to_latex
from latexword.docx.write import add_inline_latex as add_inline_math_latex
from latexword.docx.write import convert_latex_to_docx


def test_native_coloured_highlight_is_scoped_and_round_trips(tmp_path):
    document = docx.Document()
    paragraph = document.add_paragraph()
    warnings = []
    add_inline_math_latex(
        paragraph,
        r"{\sethlcolor{red}\hl{marked}}",
        warnings=warnings,
    )
    assert warnings == []
    assert paragraph.runs[0].font.highlight_color == WD_COLOR_INDEX.RED

    source = tmp_path / "input"
    document.save(source)
    tex_path = tmp_path / "shadow.tex"
    tex, warnings = docx_to_latex(str(source), str(tex_path))
    assert warnings == []
    assert r"{\sethlcolor{red}\hl{marked}}" in tex
    assert r"\hlight" not in tex
    assert r"\usepackage{soul}" in tex

    tex_path.write_text(tex, encoding="utf-8")
    output = tmp_path / "output"
    _, warnings = convert_latex_to_docx(str(tex_path), str(output))
    assert warnings == []
    assert docx.Document(output).paragraphs[0].runs[0].font.highlight_color == (
        WD_COLOR_INDEX.RED
    )


def test_yellow_highlight_uses_plain_hl(tmp_path):
    document = docx.Document()
    run = document.add_paragraph().add_run("yellow")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    source = tmp_path / "input"
    document.save(source)

    tex, warnings = docx_to_latex(str(source), str(tmp_path / "shadow.tex"))
    assert warnings == []
    assert r"\hl{yellow}" in tex
    assert r"\sethlcolor" not in tex


def test_math_displaystyle_and_highlight_render_as_omml(tmp_path):
    document = docx.Document()
    paragraph = document.add_paragraph()
    warnings = []
    add_inline_math_latex(
        paragraph,
        r"\hl{$\displaystyle x^2+1$}",
        warnings=warnings,
    )
    assert warnings == []
    assert paragraph._p.find(qn("m:oMathPara")) is not None
    assert paragraph._p.find(".//" + qn("w:highlight")).get(qn("w:val")) == "yellow"


def test_highlighted_display_zone_renders_as_omml(tmp_path):
    document = docx.Document()
    paragraph = document.add_paragraph()
    warnings = []
    add_inline_math_latex(
        paragraph,
        r"\hl{\[x^2+1\]}",
        warnings=warnings,
    )
    assert warnings == []
    assert paragraph._p.find(qn("m:oMathPara")) is not None


def test_math_delimiter_inside_text_color_renders_as_omml(tmp_path):
    document = docx.Document()
    paragraph = document.add_paragraph()
    warnings = []
    add_inline_math_latex(
        paragraph,
        r"\textcolor{blue}{\(\displaystyle \tau=\tau_0\exp\left(\frac{K}{T}\right)\)}",
        warnings=warnings,
    )
    assert warnings == []
    math = paragraph._p.find(".//" + qn("m:oMath"))
    assert math is not None
    assert math.find(".//" + qn("w:color")).get(qn("w:val")) == "0000FF"
