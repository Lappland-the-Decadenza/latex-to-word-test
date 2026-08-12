"""Measured corpus defect -- ``mc:AlternateContent`` must not be read twice.

Word writes ``mc:AlternateContent`` pairs when a document is saved with
compatibility options: the live ``mc:Choice`` branch (what Word renders)
and a ``mc:Fallback`` duplicate for readers that do not understand the
Choice's format. Iterating the run subtree read *both* branches, so a
"Рис. 1" caption text box (a ``wps:wsp`` text box in Choice, a ``v:pict``
text box in Fallback) came back twice -- "Рис. 1Рис. 1" -- in the
round-tripped document (measured on two corpus documents). The Fallback
must be pruned, in the converter and in the fidelity reader alike.
"""

import os
import sys
import tempfile
import zipfile

import docx
from lxml import etree

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from latexword.docx import read as docx_read
from latexword.docx import write as latex2word

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
MC = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"


def _insert_alternate_content(path):
    """Add an mc:AlternateContent (text in both branches) to the document's
    first run, returning the new path."""
    out_path = path + ".b"
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(out_path, "w") as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name == "word/document.xml":
                root = etree.fromstring(data)
                run = root.find(".//" + W + "r")
                ac = etree.SubElement(run, MC + "AlternateContent")
                choice = etree.SubElement(ac, MC + "Choice")
                choice.set(MC + "Requires", "wps")
                r1 = etree.SubElement(choice, W + "r")
                t1 = etree.SubElement(r1, W + "t")
                t1.text = "live"
                fb = etree.SubElement(ac, MC + "Fallback")
                r2 = etree.SubElement(fb, W + "r")
                t2 = etree.SubElement(r2, W + "t")
                t2.text = "ghost"
                data = etree.tostring(root, xml_declaration=True,
                                      encoding="UTF-8", standalone=True)
            zout.writestr(name, data)
    return out_path


def test_alternate_content_fallback_is_not_doubled():
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), "plain text")
    path = os.path.join(tempfile.mkdtemp(), "ac.docx")
    d.save(path)
    hacked = _insert_alternate_content(path)
    tex, warnings = docx_read.docx_to_latex(hacked)
    assert "live" in tex
    assert "ghost" not in tex  # the Fallback branch is a compatibility duplicate
    assert "plain text" in tex
    assert warnings == []


def test_fidelity_reader_prunes_fallback_too():
    # The converter and the fidelity reader must read the same text: if the
    # instrument counted the Fallback, the fixed converter's output would
    # look like text was lost against the original (the original's XML has
    # the duplicate, the round trip has one copy).
    import docfidelity
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), "plain text")
    path = os.path.join(tempfile.mkdtemp(), "ac.docx")
    d.save(path)
    hacked = _insert_alternate_content(path)
    with zipfile.ZipFile(hacked) as z:
        root = etree.fromstring(z.read("word/document.xml"))
    run = root.find(".//" + W + "r")
    assert docfidelity._run_text(run) == "plain textlive"
