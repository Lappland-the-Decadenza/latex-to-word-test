import json
import zipfile

import pytest
from docx import Document as WordDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from latexword.cli.workspace import main
from latexword.session.commands import workspace_apply, workspace_check, workspace_diff
from latexword.session.editdir import EditDirError, collect, create_edit_dir
from latexword.workspace.create import create_workspace
from latexword.workspace.block_render import _preserve_style


def _source(path, text="Hello world"):
    document = WordDocument()
    document.add_paragraph(text)
    document.save(path)


def test_edit_directory_contains_only_safe_surface_and_collects_answers(tmp_path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    surface = create_edit_dir("shadow", root=tmp_path / "outside")

    assert workspace not in surface.path.parents
    assert {item.name for item in surface.path.iterdir()} == {"shadow.tex", "check.cmd"}
    exposed = "".join(item.read_text(encoding="utf-8") for item in surface.path.iterdir())
    assert str(workspace) not in exposed
    assert "NodeId" not in exposed and "semantic_hash" not in exposed and "byte offset" not in exposed
    assert not (surface.path / "TASK.md").exists()
    assert collect(surface).changed is False
    (surface.path / "ANSWER.md").write_text("answer", encoding="utf-8")
    assert collect(surface).answer_text == "answer"
    (surface.path / "shadow.tex").write_text("changed", encoding="utf-8")
    result = collect(surface)
    assert result.changed and result.answer_text == "answer"


def test_existing_edit_directory_is_a_clean_edit_error(tmp_path):
    root = tmp_path / "edit"
    root.mkdir()
    with pytest.raises(EditDirError, match="edit directory already exists"):
        create_edit_dir("shadow", root=root)


def test_missing_working_copy_is_named(tmp_path):
    surface = create_edit_dir("shadow", root=tmp_path / "edit")
    (surface.path / "shadow.tex").unlink()
    with pytest.raises(EditDirError, match="working-copy-missing"):
        collect(surface)


@pytest.mark.parametrize("reference, message", [
    ("../outside.png", "unsafe-resource-path"),
    ("payload.svg", "unsupported-image-type"),
])
def test_referenced_resource_boundary_is_explicit(tmp_path, reference, message):
    surface = create_edit_dir("shadow", root=tmp_path / "edit")
    if reference == "payload.svg":
        (surface.path / reference).write_text("<svg/>", encoding="utf-8")
    (surface.path / "shadow.tex").write_text(
        rf"\includegraphics{{{reference}}}", encoding="utf-8",
    )

    with pytest.raises(EditDirError, match=message):
        collect(surface)


def test_file_checker_reports_unknown_command_location(tmp_path):
    shadow = tmp_path / "shadow.tex"
    shadow.write_text(
        "\\documentclass{article}\n\\begin{document}\n\\notarealcommand\n\\end{document}\n",
        encoding="utf-8",
    )
    result = workspace_check(shadow)
    assert not result.valid
    assert any(item["code"] == "unknown-command" for item in result.diagnostics)
    diagnostic = next(item for item in result.diagnostics if item["code"] == "unknown-command")
    assert diagnostic["line"] == 3 and diagnostic["column"] == 1


def test_file_checker_does_not_reject_document_class(tmp_path):
    shadow = tmp_path / "shadow.tex"
    shadow.write_text(
        "\\documentclass{book}\n\\begin{document}\nText\n\\end{document}\n",
        encoding="utf-8",
    )
    result = workspace_check(shadow)
    assert result.valid
    assert not any(item["code"] == "unapproved-package" for item in result.diagnostics)


def test_file_checker_accepts_the_converter_unicode_envelope(tmp_path):
    shadow = tmp_path / "shadow.tex"
    shadow.write_text(
        r"""\documentclass{article}
\usepackage{amsmath,amssymb,amsfonts,mathtools,fontspec,endnotes}
\usepackage{soul,xcolor}
\setmainfont{DejaVu Serif}
\setmonofont{DejaVu Sans Mono}
\makeatletter
\def\lTwoWHighlightColor{yellow}
\renewcommand{\sethlcolor}[1]{\def\lTwoWHighlightColor{#1}}
\renewcommand{\hl}[1]{\colorbox{\lTwoWHighlightColor}{#1}}
\makeatother
\begin{document}
Привет \hl{мир}\endnote{заметка}
\theendnotes
\end{document}
""", encoding="utf-8",
    )
    result = workspace_check(shadow)
    assert result.valid, result.diagnostics


def test_answer_and_stray_files_produce_no_operations_and_identical_output(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    workspace = tmp_path / "workspace"
    _source(source)
    create_workspace(source, workspace)
    edit = workspace
    assert main(["--quiet", "turn-start", str(workspace)]) == 0
    (edit / "ANSWER.md").write_text("response", encoding="utf-8")
    (edit / "stray.py").write_text("raise RuntimeError", encoding="utf-8")
    (edit / "TASK.md").write_text("unrelated", encoding="utf-8")

    edit_result, reconciliation = workspace_diff(workspace)
    assert edit_result.answer_text == "response" and reconciliation is None
    result = workspace_apply(workspace, output)
    assert result["changed_blocks"] == 0 and result["answer"] == "response"
    assert source.read_bytes() == output.read_bytes()


def test_invalid_shadow_reports_location_and_publishes_nothing(tmp_path, capsys):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    workspace = tmp_path / "workspace"
    _source(source)
    create_workspace(source, workspace)
    edit = workspace
    assert main(["--quiet", "turn-start", str(workspace)]) == 0
    (edit / "shadow.tex").write_text("\\begin{document}", encoding="utf-8")

    assert main(["--quiet", "workspace-apply", str(workspace), "--output", str(output)]) == 2
    assert not output.exists()
    error = capsys.readouterr().err
    assert "line=" in error and "column=" in error


def test_cli_json_is_single_envelope(tmp_path, capsys):
    shadow = tmp_path / "shadow.tex"
    shadow.write_text(
        "\\documentclass{article}\n\\begin{document}\nText\n\\end{document}\n",
        encoding="utf-8",
    )
    assert main(["--quiet", "workspace-check", str(shadow), "--json"]) == 0
    lines = capsys.readouterr().out.splitlines()
    assert len(lines) == 1
    assert json.loads(lines[0])["schema"] == "lw-cli/v1"


def test_document_commands_create_reuse_and_apply_one_managed_folder(tmp_path):
    source = tmp_path / "report.docx"
    _source(source)

    assert main(["--quiet", "edit-start", str(source)]) == 0
    managed = tmp_path / "report.latexword"
    assert managed.is_dir()
    assert (managed / "shadow.tex").is_file()
    assert not (managed / "TASK.md").exists()
    assert (managed / ".service" / "current.docx").is_file()
    assert (managed / ".service" / "original.docx").is_file()
    assert not (managed / "current.docx").exists()

    shadow = managed / "shadow.tex"
    shadow.write_text(shadow.read_text(encoding="utf-8").replace("world", "there"), encoding="utf-8")
    assert main(["--quiet", "edit-apply", str(source)]) == 0
    assert "Hello there" in "\n".join(p.text for p in WordDocument(source).paragraphs)
    assert "Hello world" in "\n".join(
        p.text for p in WordDocument(managed / ".service" / "original.docx").paragraphs
    )
    assert (managed / ".service" / "diagnostics.log").is_file()
    assert not (tmp_path / "report.figures").exists()
    assert not (tmp_path / "report_reversed.objects").exists()
    assert not (tmp_path / "report_reversed.tex").exists()

    assert main(["--quiet", "edit-start", str(source)]) == 0
    shadow.write_text(shadow.read_text(encoding="utf-8").replace("there", "again"), encoding="utf-8")
    assert main(["--quiet", "edit-apply", str(source)]) == 0
    assert "Hello again" in "\n".join(p.text for p in WordDocument(source).paragraphs)


def test_formatting_edit_reconciles_without_a_nested_agent(tmp_path):
    source = tmp_path / "source.docx"
    _source(source)
    assert main(["--quiet", "edit-start", str(source)]) == 0
    shadow = tmp_path / "source.latexword" / "shadow.tex"
    shadow.write_text(
        shadow.read_text(encoding="utf-8").replace("world", r"\textbf{world}"),
        encoding="utf-8",
    )
    assert main(["--quiet", "edit-apply", str(source)]) == 0
    paragraph = WordDocument(source).paragraphs[0]
    assert paragraph.text == "Hello world"
    assert any(run.text == "world" and run.bold for run in paragraph.runs)


def test_changed_heading_keeps_non_plain_paragraph_style(tmp_path):
    source = tmp_path / "source.docx"
    document = WordDocument()
    document.styles.add_style("CustomHeading", WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph("world", style="CustomHeading")
    document.save(source)
    assert main(["--quiet", "edit-start", str(source)]) == 0
    shadow = tmp_path / "source.latexword" / "shadow.tex"
    shadow.write_text(
        shadow.read_text(encoding="utf-8").replace("world", r"\section{Changed}"),
        encoding="utf-8",
    )
    assert main(["--quiet", "edit-apply", str(source)]) == 0
    assert WordDocument(source).paragraphs[0].style.style_id == "CustomHeading"


def test_changed_block_keeps_source_language_for_spellcheck(tmp_path):
    source = tmp_path / "source.docx"
    document = WordDocument()
    run = document.add_paragraph().add_run("Hello world")
    language = OxmlElement("w:lang")
    language.set(qn("w:val"), "uk-UA")
    run._r.get_or_add_rPr().append(language)
    document.save(source)

    assert main(["--quiet", "edit-start", str(source)]) == 0
    shadow = tmp_path / "source.latexword" / "shadow.tex"
    shadow.write_text(
        shadow.read_text(encoding="utf-8").replace("Hello world", "Changed"),
        encoding="utf-8",
    )
    assert main(["--quiet", "edit-apply", str(source)]) == 0

    with zipfile.ZipFile(source) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert 'w:lang w:val="uk-UA"' in xml


def test_changed_math_restores_source_double_ascii_spaces(tmp_path):
    template = tmp_path / "source.docx"
    _source(template)
    payload = (
        b'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        b'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        b'<m:oMath><m:r><m:t>\xe2\x80\x83\xe2\x80\x83</m:t></m:r></m:oMath></w:p>'
    )
    rendered = _preserve_style(
        (payload,), template, None, spacing_hint="word-double-space"
    )
    root = etree.fromstring(rendered[0])
    text = root.xpath("string(.//m:t)", namespaces={"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"})
    assert text == "  "


def test_document_folder_collision_gets_a_suffix_without_overwrite(tmp_path):
    source = tmp_path / "report.docx"
    _source(source)
    collision = tmp_path / "report.latexword"
    collision.mkdir()
    (collision / "keep.txt").write_text("unrelated", encoding="utf-8")

    assert main(["--quiet", "edit-start", str(source)]) == 0
    assert (collision / "keep.txt").read_text(encoding="utf-8") == "unrelated"
    assert (tmp_path / "report.latexword-2" / "shadow.tex").is_file()


def test_owned_folder_rebuilds_when_the_source_docx_is_stale(tmp_path):
    source = tmp_path / "report.docx"
    _source(source, "Before")
    assert main(["--quiet", "edit-start", str(source)]) == 0
    managed = tmp_path / "report.latexword"
    assert main(["--quiet", "edit-apply", str(source)]) == 0

    _source(source, "Changed outside the tool")
    assert main(["--quiet", "edit-start", str(source)]) == 0
    shadow = managed.joinpath("shadow.tex").read_text(encoding="utf-8")
    assert "Changed outside the tool" in shadow
    assert "Before" in "\n".join(
        p.text for p in WordDocument(managed / ".service" / "original.docx").paragraphs
    )


def test_legacy_base_workspace_is_migrated_to_original_and_current(tmp_path):
    source = tmp_path / "report.docx"
    workspace = tmp_path / "report.latexword"
    _source(source, "First")
    create_workspace(source, workspace)
    service = workspace / ".service"
    (service / "current.docx").rename(service / "base.docx")
    (service / "original.docx").unlink()

    assert main(["--quiet", "edit-start", str(source)]) == 0
    assert (service / "current.docx").is_file()
    assert (service / "original.docx").is_file()
    assert "First" in "\n".join(p.text for p in WordDocument(service / "original.docx").paragraphs)


def test_open_word_candidate_is_retained_until_the_document_closes(tmp_path, monkeypatch):
    source = tmp_path / "report.docx"
    _source(source)
    assert main(["--quiet", "edit-start", str(source)]) == 0
    managed = tmp_path / "report.latexword"
    shadow = managed / "shadow.tex"
    shadow.write_text(shadow.read_text(encoding="utf-8").replace("world", "there"), encoding="utf-8")

    monkeypatch.setattr("latexword.session.publication.word_document_is_open", lambda _path: True)
    assert main(["--quiet", "edit-apply", str(source)]) == 0
    assert "Hello world" in "\n".join(p.text for p in WordDocument(source).paragraphs)
    assert (managed / ".service" / "candidate.docx").is_file()

    monkeypatch.setattr("latexword.session.publication.word_document_is_open", lambda _path: False)
    assert main(["--quiet", "edit-apply", str(source)]) == 0
    assert "Hello there" in "\n".join(p.text for p in WordDocument(source).paragraphs)
