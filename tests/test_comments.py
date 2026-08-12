"""PLAN.md §7.1 -- comments (comments.xml), both directions.

Before this change, ``word/comments.xml`` was never read and ``\\todo`` was
an unknown command: comment text vanished entirely on a round trip, and the
plan names comments a high-value item for Plan 2 (a comment is literally an
instruction to the model).

These tests exercise the forward and reverse halves independently plus true
round trips through the real pipeline (the same shape as the smoke that
drove the implementation), rather than depending on the private corpus
documents the fidelity report measures against -- the corpus carries no
comments at all, so every shape here is synthetic.
"""

import os
import sys
import tempfile
import zipfile

import docx
from lxml import etree

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from latexword.docx import read as docx_read
from latexword.docx import write as latex2word

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


def _comments_defs(path):
    """Comment definitions of one part: id -> (author, date, text)."""
    with zipfile.ZipFile(path) as z:
        if "word/comments.xml" not in z.namelist():
            return {}
        root = etree.fromstring(z.read("word/comments.xml"))
    out = {}
    for el in root.iter(W + "comment"):
        cid = el.get(W + "id")
        if cid is not None:
            out[cid] = (el.get(W + "author") or "", el.get(W + "date") or "",
                        "".join(t.text or "" for t in el.iter(W + "t")))
    return out


def _build(text, warnings=None):
    """The real pipeline: .tex file -> .docx.

    Accepts a body fragment or a complete document (the reverse's output is
    a complete document; wrapping it again would leak its preamble into the
    body of the next generation)."""
    if "\\documentclass" not in text:
        text = ("\\documentclass{article}\n"
                "\\begin{document}\n" + text + "\n"
                "\\end{document}\n")
    tmp = tempfile.mkdtemp()
    tp = os.path.join(tmp, "src.tex")
    with open(tp, "w", encoding="utf-8") as fh:
        fh.write(text)
    out = os.path.join(tmp, "g1.docx")
    if warnings is not None:
        warnings.extend(latex2word.convert_latex_to_docx(tp, out)[1])
    else:
        latex2word.convert_latex_to_docx(tp, out)
    return out


# --- Forward: standard todo creates a real comments part -------------------


def test_comment_wires_part_content_type_and_relationship():
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), r"a\todo[inline]{note} b")
    path = os.path.join(tempfile.mkdtemp(), "c.docx")
    d.save(path)
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        ct = z.read("[Content_Types].xml").decode("utf-8")
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        docroot = etree.fromstring(z.read("word/document.xml"))
    assert "word/comments.xml" in names
    assert "comments+xml" in ct
    assert 'Target="comments.xml"' in rels
    # point comment: start and end adjacent, one reference run, all one id
    start = docroot.find(".//" + W + "commentRangeStart")
    end = docroot.find(".//" + W + "commentRangeEnd")
    refs = list(docroot.iter(W + "commentReference"))
    assert len(refs) == 1
    assert start.get(W + "id") == end.get(W + "id") == refs[0].get(W + "id")
    assert end.getprevious() is start  # nothing between: a point, not a range
    defs = _comments_defs(path)
    assert len(defs) == 1
    cid = start.get(W + "id")
    assert defs[cid][:2] == ("", "")
    assert "note" in defs[cid][2]


def test_comment_body_survives_with_formatting_and_math():
    d = docx.Document()
    latex2word.add_inline_latex(
        d.add_paragraph(), r"a\todo[inline]{see \textbf{note} and $x^2$} b")
    path = os.path.join(tempfile.mkdtemp(), "c.docx")
    d.save(path)
    defs = _comments_defs(path)
    assert len(defs) == 1
    text = next(iter(defs.values()))[2]
    assert "see note and" in text  # \textbf is formatting; text survives
    # the oMath object is inside the comments part, not dropped
    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read("word/comments.xml"))
    assert root.find(".//" + M + "oMath") is not None


def test_hyperlink_in_comment_body_keeps_its_namespace():
    # The defect that hit footnotes: an lxml tree whose root declared no
    # namespaces gets invented prefixes at serialization, and the r:id
    # attribute collides with the generated prefix -- the serializer
    # rebinds the prefix on the hyperlink element itself, changing its tag
    # to the relationships namespace. The explicit nsmap on the comments
    # root must prevent the same silent loss here.
    d = docx.Document()
    latex2word.add_inline_latex(
        d.add_paragraph(), r"a\todo[inline]{see \href{https://example.com}{link}} b")
    path = os.path.join(tempfile.mkdtemp(), "c.docx")
    d.save(path)
    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read("word/comments.xml"))
    hlinks = root.iter(W + "hyperlink")
    assert sum(1 for _ in hlinks) == 1
    assert "link" in next(iter(_comments_defs(path).values()))[2]


def test_forward_short_todo_is_named():
    # A standard todo carries only its visible body in native LaTeX.
    d = docx.Document()
    w = []
    latex2word.add_inline_latex(
        d.add_paragraph(), r"a\todo[inline]{fix this} b", warnings=w)
    path = os.path.join(tempfile.mkdtemp(), "c.docx")
    d.save(path)
    assert w == []
    defs = _comments_defs(path)
    assert len(defs) == 1
    author, date, text = next(iter(defs.values()))
    assert author == "" and date == ""
    assert "fix this" in text


def test_no_comments_writes_no_part():
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), "plain text")
    path = os.path.join(tempfile.mkdtemp(), "c.docx")
    d.save(path)
    with zipfile.ZipFile(path) as z:
        assert "word/comments.xml" not in z.namelist()


# --- Reverse: comments.xml becomes the standard todo form ------------------


def test_reverse_emits_todo_with_metadata():
    d = docx.Document()
    latex2word.add_inline_latex(
        d.add_paragraph(), r"before \todo[inline]{see \textbf{note} and $x^{2}$} after")
    path = os.path.join(tempfile.mkdtemp(), "c.docx")
    d.save(path)
    tex, warnings = docx_read.docx_to_latex(path)
    assert r"\todo[inline]{see \textbf{note} and $x^{2}$}" in tex
    assert warnings == []  # a point comment round-trips quietly


def test_reverse_range_comment_is_named():
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), r"a\todo[inline]{note} b")
    path = os.path.join(tempfile.mkdtemp(), "c.docx")
    d.save(path)
    # give the range a body: insert a text run between start and end
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(path + ".b", "w") as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name == "word/document.xml":
                root = etree.fromstring(data)
                start = root.find(".//" + W + "commentRangeStart")
                r = etree.Element(W + "r")
                t = etree.SubElement(r, W + "t")
                t.text = "XX"
                start.addnext(r)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(name, data)
    tex, warnings = docx_read.docx_to_latex(path + ".b")
    assert r"\todo[inline]{note}" in tex
    assert "XX" in tex  # the covered text is preserved, only the highlight is lost
    assert any("covers text" in w for w in warnings)


def test_reverse_spanning_comment_is_named():
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), r"a\todo[inline]{note} b")
    d.add_paragraph("next paragraph")
    path = os.path.join(tempfile.mkdtemp(), "c.docx")
    d.save(path)
    # move the range end into the next paragraph
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(path + ".b", "w") as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name == "word/document.xml":
                root = etree.fromstring(data)
                paras = root.findall(".//" + W + "p")
                end = root.find(".//" + W + "commentRangeEnd")
                paras[1].append(end)  # moves the element
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(name, data)
    tex, warnings = docx_read.docx_to_latex(path + ".b")
    assert r"\todo[inline]{note}" in tex
    assert "next paragraph" in tex  # the covered text is preserved
    assert any("spans paragraphs" in w for w in warnings)


def test_reverse_reply_is_named():
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), r"a\todo[inline]{note} b")
    path = os.path.join(tempfile.mkdtemp(), "c.docx")
    d.save(path)
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(path + ".b", "w") as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name == "word/comments.xml":
                root = etree.fromstring(data)
                root.find(W + "comment").set(W + "parent", "1")
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(name, data)
    tex, warnings = docx_read.docx_to_latex(path + ".b")
    assert r"\todo[inline]{note}" in tex  # the reply's text is carried
    assert any("reply" in w for w in warnings)  # the threading loss is named


def test_reverse_dangling_reference_is_named():
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), r"t\todo[inline]{real}")
    path = os.path.join(tempfile.mkdtemp(), "c.docx")
    d.save(path)
    # break the document: point the reference at a nonexistent id
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(path + ".b", "w") as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name == "word/document.xml":
                root = etree.fromstring(data)
                ref = root.find(".//" + W + "commentReference")
                ref.set(W + "id", "99")
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(name, data)
    tex, warnings = docx_read.docx_to_latex(path + ".b")
    assert r"\todo[inline]" not in tex
    assert any("no definition" in w for w in warnings)


def test_reverse_unreferenced_definition_is_named():
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), r"a\todo[inline]{note} b")
    path = os.path.join(tempfile.mkdtemp(), "c.docx")
    d.save(path)
    # add a second definition no reference points at
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(path + ".b", "w") as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name == "word/comments.xml":
                root = etree.fromstring(data)
                c = etree.SubElement(root, W + "comment")
                c.set(W + "id", "2")
                c.set(W + "author", "Orphan")
                p = etree.SubElement(c, W + "p")
                r = etree.SubElement(p, W + "r")
                t = etree.SubElement(r, W + "t")
                t.text = "unused"
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(name, data)
    tex, warnings = docx_read.docx_to_latex(path + ".b")
    assert r"\todo[inline]{note}" in tex
    assert any("unreferenced comment id=2" in w for w in warnings)


def test_reverse_multiparagraph_comment_is_named():
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), r"a\todo[inline]{one two} b")
    path = os.path.join(tempfile.mkdtemp(), "c.docx")
    d.save(path)
    # split the comment's paragraph into two paragraphs (move the last run)
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(path + ".b", "w") as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name == "word/comments.xml":
                root = etree.fromstring(data)
                comment = root.find(W + "comment")
                p = comment.find(W + "p")
                runs = p.findall(W + "r")
                p2 = etree.Element(W + "p")
                p2.append(runs[-1])  # moves the run: now two paragraphs
                p.addnext(p2)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(name, data)
    tex, warnings = docx_read.docx_to_latex(path + ".b")
    assert r"\todo[inline]{" in tex and "one" in tex and "two" in tex
    assert any("paragraph" in w for w in warnings)


def test_comment_in_table_cell_is_carried():
    # A comment anchored inside a table cell: the note/comment contexts
    # thread into the cell pass, so the mark is not silently skipped.
    out1 = _build("\\begin{tabular}{cc}\nA & B \\\\\nC & \\todo[inline]{in cell} \\\\\n\\end{tabular}")
    tex1, w1 = docx_read.docx_to_latex(out1)
    assert r"\todo[inline]{in cell}" in tex1
    assert w1 == []


# --- Round trip through the real pipeline -----------------------------------


def test_comment_round_trip_is_stable():
    src = r"Body\todo[inline]{note one} more \todo[inline]{see \textbf{note} and $x^2$}."
    out1 = _build(src)
    tex1, w1 = docx_read.docx_to_latex(out1)
    assert r"\todo[inline]{note one}" in tex1
    assert r"\todo[inline]{see \textbf{note} and $x^{2}$}" in tex1
    assert "see \\textbf{note} and $x^{2}$" in tex1
    assert w1 == []
    out2 = _build(tex1)
    tex2, w2 = docx_read.docx_to_latex(out2)
    assert tex2 == tex1  # fixed point from generation 2
    assert w2 == []


def test_todo_preamble_emitted_only_when_needed():
    # no comments -> no \todo definition in the preamble, byte-identical
    d = docx.Document()
    latex2word.add_inline_latex(d.add_paragraph(), "plain text")
    path = os.path.join(tempfile.mkdtemp(), "c.docx")
    d.save(path)
    tex, w = docx_read.docx_to_latex(path)
    assert "\\usepackage{todonotes}" not in tex
    assert w == []
    # comments -> the definition is present, and the reverse output must
    # convert back without the definition leaking into the body or
    # producing warnings
    d2 = docx.Document()
    latex2word.add_inline_latex(d2.add_paragraph(), r"a\todo[inline]{n} b")
    path2 = os.path.join(tempfile.mkdtemp(), "c2.docx")
    d2.save(path2)
    tex2, w2 = docx_read.docx_to_latex(path2)
    assert "\\usepackage{todonotes}" in tex2
    assert w2 == []
    tmp = tempfile.mkdtemp()
    tp = os.path.join(tmp, "t.tex")
    with open(tp, "w", encoding="utf-8") as fh:
        fh.write(tex2)
    _, w3 = latex2word.convert_latex_to_docx(tp, os.path.join(tmp, "t.docx"))
    assert w3 == []
