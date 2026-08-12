"""Named Word styles travel through the detached sidecar, not native TeX."""

import json
import os

import docx

from latexword.docx import read as docx_read
from latexword.docx import write as latex2word


def _build_docx_with_style(tmp_path, style_id, text):
    document = docx.Document()
    document.styles.add_style(style_id, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    paragraph = document.add_paragraph(text)
    paragraph._p.style = style_id
    path = tmp_path / "source.docx"
    document.save(path)
    return path


def test_named_style_is_detached_and_restored(tmp_path):
    source = _build_docx_with_style(tmp_path, "CustomAbstract", "Body")
    tex_path = tmp_path / "shadow.tex"
    tex, warnings = docx_read.docx_to_latex(str(source), str(tex_path))

    assert warnings == []
    assert "Body" in tex
    assert "CustomAbstract" not in tex
    manifest = json.loads(
        (tmp_path / "shadow.objects" / "manifest.json").read_text(
            encoding="utf-8"
        )
    )
    assert any(item["kind"] == "paragraph-style"
               for item in manifest["attachments"])

    tex_path.write_text(tex, encoding="utf-8")
    output = tmp_path / "restored.docx"
    _, warnings = latex2word.convert_latex_to_docx(
        str(tex_path), str(output), str(source), reference_mode="copy"
    )
    assert warnings == []
    assert docx.Document(output).paragraphs[0].style.style_id == (
        "CustomAbstract"
    )


def test_native_paragraphs_do_not_synthesize_named_styles(tmp_path):
    tex_path = tmp_path / "input.tex"
    tex_path.write_text(
        "\\documentclass{article}\n\\begin{document}\nBody\n"
        "\\end{document}\n",
        encoding="utf-8",
    )
    output = tmp_path / "output.docx"
    _, warnings = latex2word.convert_latex_to_docx(str(tex_path), str(output))
    assert warnings == []
    assert docx.Document(output).paragraphs[0].text.strip() == "Body"


def test_named_style_does_not_enter_the_generated_preamble(tmp_path):
    source = _build_docx_with_style(tmp_path, "CustomAbstract", "Body")
    tex, warnings = docx_read.docx_to_latex(str(source), str(tmp_path / "out.tex"))
    assert warnings == []
    assert "newenvironment" not in tex
    assert "newcommand" not in tex


def test_section_named_styles_become_native_headings(tmp_path):
    document = docx.Document()
    section = document.styles.add_style(
        "JnepSection", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH
    )
    subsection = document.styles.add_style(
        "JnepSubsection", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH
    )
    document.add_paragraph("Main section")._p.style = section.style_id
    document.add_paragraph("Nested section")._p.style = subsection.style_id
    source = tmp_path / "source.docx"
    document.save(source)

    tex_path = tmp_path / "shadow.tex"
    tex, warnings = docx_read.docx_to_latex(str(source), str(tex_path))
    assert warnings == []
    assert r"\section{Main section}" in tex
    assert r"\subsection{Nested section}" in tex

    tex_path.write_text(tex, encoding="utf-8")
    output = tmp_path / "restored.docx"
    _, warnings = latex2word.convert_latex_to_docx(
        str(tex_path), str(output), str(source), reference_mode="copy"
    )
    assert warnings == []
    paragraphs = docx.Document(output).paragraphs
    assert [p.style.style_id for p in paragraphs[:2]] == [
        "JnepSection", "JnepSubsection"
    ]
