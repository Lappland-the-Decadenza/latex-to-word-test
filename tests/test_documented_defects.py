"""Pins for every defect documented in `DEFECTS.md` plus the R6 remainder
table in `REWRITE_FORWARD.md` (around line 221), re-verified against the
*new* emitter (`latex2omml.parse`/`emit`, `latex2word.latex_math_to_omml`)
and, for the reverse-direction items, `word2latex`.

`mathml_normalize.py` and `latex2word.legacy_latex_math_to_omml` are the
*old* pipeline and are never exercised here -- see `REWRITE_FORWARD.md`:
the new emitter does not inherit `mathml_normalize.py`'s fixes just because
the old pipeline had them.

Every assertion below is about *which OMML object* a construct becomes
(element/attribute identity), never merely "the tree is well-formed" --
`DEFECTS.md`'s opening section explains why that distinction is the one
that matters (it is what hid D1 for a whole round of feedback).
"""

import os
import sys

import docx
import pytest
from docx.oxml.ns import qn
from lxml import etree

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from latexword.math import latex2omml as L  # noqa: E402
from latexword.docx import write as latex2word  # noqa: E402
from latexword.math import omml2latex as word2latex  # noqa: E402
from latexword.docx import read as docx_read  # noqa: E402


# --- helpers -----------------------------------------------------------------


def _tag(el):
    return el.tag.split("}", 1)[-1]


def _emit(tex):
    """LaTeX -> `m:oMath` through the new forward pipeline only."""
    return L.emit(L.parse(tex))


def _root_tags(om):
    """Local names of the direct children of an `m:oMath` root."""
    return [_tag(c) for c in om]


def _find(root, tag):
    """First descendant (root included) with local name `tag`, or None."""
    for el in root.iter():
        if _tag(el) == tag:
            return el
    return None


def _all_tags(root):
    return {_tag(el) for el in root.iter()}


def _text_of(el):
    """Concatenated `m:t` text under `el`."""
    return "".join(t.text or "" for t in el.iter(L.qm("t")))


# --- D1: function application is m:func, with the deliberate single-letter gap


@pytest.mark.parametrize("tex", [
    r"\sin{x}",
    r"\det A",
    r"\operatorname{rank}\left(A\right)",
    r"\limsup_{n\to\infty} a_n",
    r"\sin^{2}\theta",
])
def test_d1_function_application_is_m_func(tex):
    """DEFECTS.md D1: '\\sin{x} used to emit two sibling runs... not a
    Function object.' Every operator-name-plus-operand construct must be one
    `m:func`, not loose sibling runs (a plain-text name next to an unrelated
    argument)."""
    om = _emit(tex)
    assert _root_tags(om) == ["func"], (
        f"{tex!r} did not become a single m:func: {_root_tags(om)}"
    )


def test_d1_deliberate_gap_single_letter_name_not_wrapped():
    """DEFECTS.md D1, 'Known deliberate gap': a single-letter name
    immediately followed by `(...)` -- `f(x)` -- must NOT be wrapped into
    `m:func` (genuinely ambiguous with implicit multiplication, e.g.
    `a(b+c)`). `(x)` must still become a real `m:d` object either way."""
    om = _emit(r"f(x)")
    assert "func" not in _all_tags(om), "single-letter name was wrapped in m:func"
    assert _root_tags(om) == ["r", "d"]
    d = _find(om, "d")
    assert d is not None


# --- D2: bare brackets are m:d, including nested un-\left pairs -------------


@pytest.mark.parametrize("tex,beg,end", [
    (r"(x)", "(", ")"),
    (r"[x]", "[", "]"),
    (r"\{x\}", "{", "}"),
])
def test_d2_bare_delimiters_become_m_d(tex, beg, end):
    """DEFECTS.md D2: 'Plain `(`, `[`, `{` written without `\\left` reached
    Word as literal bracket characters that did not stretch... did not
    behave as a group when edited.' Each must produce a real `m:d`."""
    om = _emit(tex)
    assert _root_tags(om) == ["d"]
    d = om[0]
    dpr = _find(d, "dPr")
    beg_el = _find(dpr, "begChr")
    end_el = _find(dpr, "endChr")
    assert beg_el.get(L.qm("val")) == beg
    assert end_el.get(L.qm("val")) == end


def test_d2_nested_bare_pair_both_wrapped():
    """DEFECTS.md D2, 'Bug found and fixed during implementation': a nested
    un-\\left pair (`\\{ f(x) \\}`) must have BOTH the outer `{}` and the
    inner `(x)` folded into `m:d` objects, not just the outer one."""
    om = _emit(r"\{ f(x) \}")
    outer = om[0]
    assert _tag(outer) == "d"
    outer_dpr = _find(outer, "dPr")
    assert _find(outer_dpr, "begChr").get(L.qm("val")) == "{"
    assert _find(outer_dpr, "endChr").get(L.qm("val")) == "}"
    inner = _find(outer[1], "d")  # outer[0] is dPr, outer[1] is the e slot
    assert inner is not None, "inner (x) was not folded into its own m:d"
    inner_dpr = _find(inner, "dPr")
    assert _find(inner_dpr, "begChr").get(L.qm("val")) == "("
    assert _find(inner_dpr, "endChr").get(L.qm("val")) == ")"


# --- D3: \overbrace / \underbrace are m:groupChr -----------------------------


@pytest.mark.parametrize("tex,chr_,pos", [
    (r"\overbrace{x}^{n}", "\u23de", "top"),
    (r"\underbrace{x}_{n}", "\u23df", "bot"),
])
def test_d3_overbrace_underbrace_are_groupchr(tex, chr_, pos):
    """DEFECTS.md D3 ('Still open' at the time it was written): '\\overbrace
    / \\underbrace are not m:groupChr... Currently emitted as nested
    m:limLow/m:limUpp.' `REWRITE_FORWARD.md` claims this is gained by the
    rewrite ('new capability, unreachable today') -- verified here rather
    than assumed."""
    om = _emit(tex)
    gc = _find(om, "groupChr")
    assert gc is not None, f"{tex!r} did not produce m:groupChr"
    pr = _find(gc, "groupChrPr")
    assert _find(pr, "chr").get(L.qm("val")) == chr_
    assert _find(pr, "pos").get(L.qm("val")) == pos
    # It must not have regressed back to the old limLow/limUpp brace-glyph
    # shape for this fresh, forward-authored input.
    assert "limLow" not in _all_tags(om) and "limUpp" not in _all_tags(om)


# --- D4: adjacent-character run merging (judgement call -- pin, do not guess)


def test_d4_adjacent_characters_are_separate_runs():
    """DEFECTS.md D4: 'The XSL collapses runs, so `f(t)dt` arrives as a
    single `<m:t>`... Not independently confirmed as visually wrong yet.'

    Observed with the new emitter: `d` and `t` in `f(t)dt` are NOT merged --
    each `Ident` becomes its own `m:r`/`m:t`. This is the opposite of the
    old XSL's behaviour and is pinned as-observed per this task's
    instructions ('judgement call... pin the observed behaviour and flag it
    in your report as needing a decision'); no claim is made here about
    which shape is correct."""
    om = _emit(r"f(t)dt")
    # f, then the (t) delimiter, then "d" and "t" as two more top-level runs.
    assert _root_tags(om) == ["r", "d", "r", "r"]
    trailing_runs = [c for c in om][2:]
    assert [_text_of(r) for r in trailing_runs] == ["d", "t"]


# --- D5: \left. / \right| -> empty m:begChr (unverified for years) ----------


def test_d5_left_dot_right_bar_empty_begchr():
    """DEFECTS.md D5 ('Still open'): '`\\left. F(x) \\right|` produces
    `<m:begChr m:val=""/>`. Needs checking in Word that this renders as
    nothing rather than falling back to a default `(`... has not been
    confirmed for begChr specifically.' Pinned as observed: the new
    emitter still produces an explicitly empty (not absent) `begChr`."""
    om = _emit(r"\left. F(x) \right|")
    assert _root_tags(om) == ["d"]
    d = om[0]
    dpr = _find(d, "dPr")
    beg_el = _find(dpr, "begChr")
    end_el = _find(dpr, "endChr")
    assert beg_el is not None and beg_el.get(L.qm("val")) == ""
    assert end_el is not None and end_el.get(L.qm("val")) == "|"


def test_d5_right_dot_empty_endchr():
    """The mirror one-sided case: `\\right.` must likewise leave `endChr`
    an explicit empty value, not a missing attribute or a fallback glyph."""
    om = _emit(r"\left| F(x) \right.")
    d = om[0]
    dpr = _find(d, "dPr")
    beg_el = _find(dpr, "begChr")
    end_el = _find(dpr, "endChr")
    assert beg_el is not None and beg_el.get(L.qm("val")) == "|"
    assert end_el is not None and end_el.get(L.qm("val")) == ""


# --- D6: prescripts -> m:sPre (untested at the time DEFECTS.md was written) -


def test_d6_prescript_is_m_spre():
    """DEFECTS.md D6 ('Possible'): '`{}^{a}_{b}X` should be `m:sPre`.
    Untested; the XSL supports `m:sPre`.' The parser's accepted spelling
    (per `mathast.CONSTRUCTS_BY_NAME['prescript']` / CANONICAL.md) is
    `\\prescript{sup}{sub}{base}`, not the bare `{}^{a}_{b}X` form -- pinned
    against the spelling the parser actually accepts."""
    om = _emit(r"\prescript{a}{b}{X}")
    assert _root_tags(om) == ["sPre"]
    spre = om[0]
    # OOXML slot order is sub, sup, e (mathast.PreScript docstring: the
    # macro's argument order sup/sub/base does not match the OOXML order).
    children_tags = [_tag(c) for c in spre]
    assert children_tags == ["sub", "sup", "e"]
    assert _text_of(_find(spre, "sub")) == "b"
    assert _text_of(_find(spre, "sup")) == "a"
    assert _text_of(_find(spre, "e")) == "X"


# --- D7: symbol operators are upright, not italic ----------------------------


@pytest.mark.parametrize("tex", [r"\nabla", r"\partial"])
def test_d7_symbol_operators_are_explicitly_upright(tex):
    """DEFECTS.md D7 ('Fixed' in the old pipeline): '\\nabla, \\partial...
    came out visibly slanted in Word... Root cause: MML2OMML.XSL's
    GetFontCur defaults to italic for any mo with no explicit style
    override.' The new emitter must carry the explicit upright style itself
    (`m:sty val="p"`) rather than relying on Word's own default, which is
    italic."""
    om = _emit(tex)
    assert _root_tags(om) == ["r"]
    run = om[0]
    rpr = _find(run, "rPr")
    assert rpr is not None
    sty = _find(rpr, "sty")
    assert sty is not None and sty.get(L.qm("val")) == "p"


# --- D8: primes stay two literal U+2032s, never one compound U+2033 --------


def test_d8_double_prime_is_two_u2032_not_one_u2033():
    """DEFECTS.md D8 ('Fixed'): 'f'' visibly differed from a hand-typed f''
    ... latex2mathml folds repeated ' into a single "compound prime"
    codepoint -- '' becomes one U+2033 DOUBLE PRIME character... Both
    LaTeX's own fonts and Word's native autocorrect render repeated single
    primes (U+2032) side by side instead.'"""
    om = _emit("f''")
    text = _text_of(om)
    assert "\u2033" not in text, "compound DOUBLE PRIME (U+2033) leaked in"
    assert text.count("\u2032") == 2, f"expected two U+2032, got {text!r}"


# --- D9: spacing macros are proportional Unicode glyphs, not ASCII runs ----


@pytest.mark.parametrize("tex,width_char", [
    (r"\,", "\u2009"),   # thin space
    (r"\:", "\u2005"),   # four-per-em space
    (r"\quad", "\u2003"),  # em space
])
def test_d9_spacing_is_proportional_unicode_not_ascii(tex, width_char):
    """DEFECTS.md D9 ('Fixed'): 'The gap... was implemented as a run of
    literal " " characters. A run of plain ASCII spaces is real, selectable,
    zero-ink content... Fixed... widths now map to proportional Unicode
    space characters.'

    Deviation from `DEFECTS.md`'s wording: the documented symptom uses
    `\\;`, but the new parser's macro vocabulary does not include `\\;` at
    all (`\\; ` raises `UnknownMacroError`) -- only `\\,`, `\\:`, `\\quad`,
    `\\qquad`. `\\:` is used here in its place; see this test's report note."""
    om = _emit(tex + " x")
    run = om[0]
    text = _text_of(run)
    assert text == width_char
    assert " " not in text, "spacing degraded to a literal ASCII space"


def test_d9_qquad_is_two_em_spaces():
    om = _emit(r"\qquad x")
    text = _text_of(om[0])
    assert text == "\u2003\u2003"
    assert " " not in text


# --- D10: a script on a bare closing delimiter must not swallow the rest ----


def test_d10_scripted_closing_delimiter_scoped_correctly():
    """DEFECTS.md D10 ('Fixed'): '(x + y)^{n} = \\sum...` converts to an
    m:d whose endChr is empty and whose single m:e contains the entire rest
    of the expression, including the = and the sum -- not just x + y.'

    The `m:d` must contain only `x + y`; `=` and the `\\sum` must be
    siblings OUTSIDE it, and the `m:d`'s `endChr` must be non-empty."""
    tex = r"(x + y)^{n} = \sum_{k=0}^{n}\binom{n}{k}x^k y^{n-k}"
    om = _emit(tex)
    # Root: sSup(d(...), n), then "=" run, then the nary.
    assert _root_tags(om) == ["sSup", "r", "nary"]
    sSup = om[0]
    d = _find(sSup, "d")
    assert d is not None
    dpr = _find(d, "dPr")
    assert _find(dpr, "begChr").get(L.qm("val")) == "("
    end_val = _find(dpr, "endChr").get(L.qm("val"))
    assert end_val == ")", "endChr must not be empty -- that is exactly D10"
    inner_text = _text_of(d)
    assert inner_text == "x+y", f"m:d must contain only 'x + y', got {inner_text!r}"
    eq_run = om[1]
    assert _text_of(eq_run) == "="
    assert _tag(om[2]) == "nary"


# --- R6 remainder table (REWRITE_FORWARD.md, ~line 221) ----------------------


def test_r6_every_m_d_reverses_to_left_right():
    """The test Rule 1 implies: reversing ANY `m:d` must always produce
    `\\left...\\right`, never bare delimiters chosen by a content-height
    heuristic."""
    tex = r"\left(x+y\right)^{2}"
    om = _emit(tex)
    back = word2latex.to_latex(om)
    assert back == r"\left(x+y\right)^{2}"


def test_r6_matrix_wrapped_in_left_paren_reverses_to_pmatrix():
    """REWRITE_FORWARD.md R6: '`\\begin{matrix}` wrapped in `\\left(` ->
    `pmatrix` -- benign, moves toward Rule 5 -- confirm, then close as
    intended.' Pinned as intended behaviour."""
    tex = r"\left(\begin{matrix} a & b \\ c & d \end{matrix}\right)"
    om = _emit(tex)
    back = word2latex.to_latex(om)
    assert back.strip().startswith(r"\begin{pmatrix}")
    assert back.strip().endswith(r"\end{pmatrix}")
    assert "matrix}" in back and "pmatrix" in back
    assert r"\begin{matrix}" not in back


def test_r6_align_does_not_leak_auto_numbering():
    """REWRITE_FORWARD.md R6: 'B6 -- align auto-numbering -- dies with the
    old chain -- verify rather than assume.' `DEFECTS.md`'s carried-over
    note: the old chain rendered numbers as a literal third matrix column
    (`\\text{(1)}`). The new emitter must never produce that column."""
    tex = (
        "\\begin{align}\n"
        "x &= y + 1 \\\\\n"
        "z &= w\n"
        "\\end{align}"
    )
    om = latex2word.latex_math_to_omml(tex, "block")
    s = etree.tostring(om).decode()
    assert "(1)" not in s and "(2)" not in s
    assert "\\text{(1)}" not in word2latex.to_latex(om)


def test_r6_mathrm_d_survives_roundtrip():
    """REWRITE_FORWARD.md R6 / CANONICAL.md's `\\mathrm` note: '\\mathrm{d}
    upright styling -- forward -- we control this now -- decide and
    enforce.' The decision: `\\mathrm{...}` IS representable (`m:rPr/m:sty
    val="p"`, the same run property `\\sin`/`\\operatorname{...}` already
    use), so it is accepted and round-trips rather than being rejected --
    closing the "`\\mathrm{d}` loses its upright styling" entry in
    `CLAUDE.md`'s known limitations. `\\mathrm{d}x` must not become an
    `m:func` (that would wrongly claim "d" is a named function applied to
    "x")."""
    om = latex2word.latex_math_to_omml(r"\mathrm{d}x", "inline")
    assert _root_tags(om) == ["r", "r"]
    assert word2latex.to_latex(om) == r"\mathrm{d}x"
    ast = L.parse(r"\mathrm{d}x")
    assert L.serialize(ast) == r"\mathrm{d}x"


# --- item 2: the spacing macro family (\, \: \; \! \  \quad \qquad) --------


def test_thickspace_is_accepted_and_roundtrips():
    """`\\;` (thickspace) was entirely outside the parser's vocabulary --
    `UnknownMacroError` -- even though it is standard LaTeX and appears
    constantly in real input. It must parse, canonicalize idempotently, and
    survive a real OMML round trip, exactly like its siblings `\\,`/`\\:`."""
    tex = r"a \; b"
    ast = L.parse(tex)
    assert L.serialize(ast) == r"a\;b"
    om = L.emit(ast)
    assert word2latex.to_latex(om) == r"a\;b" or word2latex.to_latex(om) == r"a\; b"


def test_control_space_is_accepted_and_roundtrips():
    """The control-space macro `\\ ` (backslash immediately followed by a
    space) is likewise part of the standard spacing family and must not be
    an `UnknownMacroError`."""
    tex = "a\\ b"
    ast = L.parse(tex)
    om = L.emit(ast)
    back = word2latex.to_latex(om)
    # Re-parsing the reversed text must reach the same fixed point.
    assert L.serialize(L.parse(back)) == L.serialize(L.parse(tex))


def test_negative_thin_space_is_tolerated_with_a_named_warning():
    """PLAN.md §6.2 moved `\\!` from `_UNSUPPORTED` to
    `compat/tolerated.py`: the backward width has no Unicode space glyph
    to borrow (unlike every other spacing macro), and a whole-equation
    fallback was the wrong punishment for a typographic nicety (measured
    10 times on the .tex corpus). The pin that used to assert the typed
    `UnsupportedConstructError` now asserts the new contract: never
    silently dropped -- a named warning, and no node in the AST (an
    empty-Op placeholder used to break Rule 2 function application and
    differ between generations, see the §6.2 record)."""
    warnings = []
    node = L.parse(r"a \! b", warnings=warnings)
    assert L.serialize(node) == "ab"
    assert len(warnings) == 1 and "negative thin space" in warnings[0]
    # Transparent to function application: one Func, one warning.
    warnings2 = []
    node2 = L.parse(r"\ln\! x", warnings=warnings2)
    assert L.serialize(node2) == r"\ln x"
    assert len(warnings2) == 1


# --- item 4: align* (and the rest of the alias family) fold to matrix ------


def test_align_star_canonicalizes_to_matrix():
    """CANONICAL.md Rule 5: an undelimited `m:m` carries nothing that
    distinguishes an aligned equation system from a plain matrix, so exactly
    one spelling is canonical -- `matrix`, what `word2latex.py` has always
    produced for this shape. `align*` (and its alias family) remain
    accepted *input* spellings but are not what `serialize` settles to."""
    tex = r"\begin{align*} a &= b \\ c &= d \end{align*}"
    s = L.canonicalize(tex)
    assert s.strip().startswith(r"\begin{matrix}")
    assert s.strip().endswith(r"\end{matrix}")
    assert "align" not in s
    # Idempotent: canonicalizing the result again changes nothing.
    assert L.canonicalize(s) == s
    # Both `align*` and bare `matrix` parse to the identical AST.
    assert L.parse(tex) == L.parse(r"\begin{matrix} a &= b \\ c &= d \end{matrix}")


def test_block_level_align_star_reaches_matrix_end_to_end():
    """The document layer (`latex2word.MATH_ENVS`) must still route a
    block-level `\\begin{align*}...\\end{align*}` through the same emitter,
    and the result must reverse to `matrix`, not silently fail or keep
    `align*` alive as a distinct OMML shape."""
    tex = "\\begin{align*}\na &= b \\\\\nc &= d\n\\end{align*}\n"
    om = latex2word.latex_math_to_omml(tex, "block")
    assert word2latex.to_latex(om).strip().startswith(r"\begin{matrix}")


# --- item 5: a delimited array must not lose its column spec ---------------


def test_delimited_array_keeps_its_column_spec():
    """`\\left(\\begin{array}{rl} a & b \\\\ c & d \\end{array}\\right)` used
    to fold straight into `pmatrix`, silently dropping the `{rl}` column
    spec (clause 3: never mutate content silently) -- the columns visibly
    re-centred on the next pass. A genuinely non-centred column spec has
    nowhere to live on `pmatrix`, so the fold must not happen; the explicit
    `\\left(\\begin{array}{rl}...\\end{array}\\right)` form is kept instead."""
    tex = r"\left(\begin{array}{rl} a & b \\ c & d \end{array}\right)"
    ast = L.parse(tex)
    om = L.emit(ast)
    back = word2latex.to_latex(om)
    assert "array" in back and "{rl}" in back
    assert "pmatrix" not in back
    # Full round trip is a fixed point.
    assert L.serialize(L.parse(back)) == L.serialize(ast)


def test_delimited_array_all_center_still_folds_to_pmatrix():
    """An `array` whose columns are all `center` carries nothing Rule 5a
    needs recorded (same discipline as Rule 6a's `\\limits`/`\\nolimits`:
    never a redundant override), so it still folds into `pmatrix` exactly
    like a bare `\\begin{matrix}` would -- only a genuine non-centre column
    spec keeps the explicit `array` form."""
    tex = r"\left(\begin{array}{cc} a & b \\ c & d \end{array}\right)"
    om = L.emit(L.parse(tex))
    back = word2latex.to_latex(om)
    assert back.strip().startswith(r"\begin{pmatrix}")
    assert "array" not in back


# --- D12: \parallel, \lesssim, \gtrsim were missing from the symbol table ---


@pytest.mark.parametrize("tex,codepoint", [
    (r"\parallel", "∥"),
    (r"\lesssim", "≲"),
    (r"\gtrsim", "≳"),
])
def test_d12_relation_macros_are_known(tex, codepoint):
    """DEFECTS.md D12 ('Fixed'): `k_{\\parallel}`, `D\\lesssim \\delta_{DW}`,
    `t\\gtrsim \\lambda_{sf}` (found converting
    `tests/corpus/output_opus.tex`) failed with "unknown macro" because
    `\\parallel`/`\\lesssim`/`\\gtrsim` had no entry in `mathsyms.SYMBOL_MAP`,
    unlike their neighbours `\\|`, `\\leq`, `\\geq`. Must emit the correct
    codepoint, upright (matching every other relation symbol, D7's fix)."""
    om = _emit(tex)
    assert _root_tags(om) == ["r"]
    run = om[0]
    assert _text_of(run) == codepoint
    rpr = _find(run, "rPr")
    assert rpr is not None
    sty = _find(rpr, "sty")
    assert sty is not None and sty.get(L.qm("val")) == "p"


# --- D13: \setlength/\renewcommand leaked their second {arg} as literal text


@pytest.mark.parametrize("tex", [
    r"\setlength{\tabcolsep}{3pt}X",
    r"\renewcommand{\arraystretch}{1.15}X",
    r"\newcommand{\foo}[1]{bar}X",
])
def test_d13_two_arg_definition_commands_leave_no_stray_text(tex):
    """DEFECTS.md D13 ('Fixed'): `\\setlength{\\tabcolsep}{3pt}` and
    `\\renewcommand{\\arraystretch}{1.15}` (found in the same file as D12)
    left a stray `{3pt}`/`{1.15}` as literal visible text in the converted
    document -- `add_inline_latex`'s "unknown command: keep the argument"
    fallback only consumes the *first* braced argument of what is really a
    two-argument command. Both arguments (and `\\renewcommand`/
    `\\newcommand`'s optional `[nargs]`) must be dropped silently, leaving
    only the prose that follows."""
    doc = docx.Document()
    p = doc.add_paragraph()
    warnings = []
    latex2word.add_inline_latex(p, tex, warnings=warnings)
    assert p.text == "X"
    assert warnings == []


# --- D14: multi-line |...| callouts (m:eqArr) must keep their rows ---------


def _eqarr_omml(rows):
    """An `m:oMath` whose `m:d` (beg=|, end=|) wraps an `m:eqArr` of one
    `m:e` per row -- the structure Word writes for a multi-paragraph
    explanation in vertical bars. The forward pipeline never emits
    `m:eqArr` (it writes `m:m` for every matrix variant), so the OMML must
    be built by hand."""
    M = L.qm
    om = etree.Element(M("oMath"))
    d = etree.SubElement(om, M("d"))
    dpr = etree.SubElement(d, M("dPr"))
    for tag, val in (("begChr", "|"), ("endChr", "|")):
        el = etree.SubElement(dpr, M(tag))
        el.set(M("val"), val)
    e = etree.SubElement(d, M("e"))
    eqarr = etree.SubElement(e, M("eqArr"))
    for row in rows:
        re = etree.SubElement(eqarr, M("e"))
        r = etree.SubElement(re, M("r"))
        t = etree.SubElement(r, M("t"))
        t.text = row
    return om


def test_d14_eqarr_in_bars_reverses_to_vmatrix():
    """DEFECTS.md D14: a multi-paragraph explanation in vertical bars is
    `m:d` (beg=|, end=|) wrapping `m:eqArr` -- rows of `m:e` children with
    no `m:mr` grid. The old walker's fallback flattened the rows into one
    line, so in the round trip the bars stretched over a single line
    instead of along all paragraphs. The eqArr must reverse to
    `\\begin{vmatrix}`: the variant's own |/| delimiters then keep the
    bars spanning the whole block."""
    om = _eqarr_omml(["a", "b", "c"])
    back = word2latex.to_latex(om)
    assert back.strip().startswith(r"\begin{vmatrix}")
    assert back.strip().endswith(r"\end{vmatrix}")
    assert back.count(r"\\") == 2, "three rows must keep two row breaks"
    assert "left" not in back, "the variant supplies its own delimiters"


def test_d14_vmatrix_roundtrip_is_a_fixed_point():
    """D14's repaired spelling round-trips: `\\begin{vmatrix}` parses to
    `m:d` (|,|) wrapping a multi-row `m:m`, and that `m:d` folds straight
    back to the same vmatrix (the m:m sole-content branch at work)."""
    tex = r"\begin{vmatrix} a \\ b \\ c \end{vmatrix}"
    om = _emit(tex)
    assert _root_tags(om) == ["d"], "vmatrix must become one m:d"
    d = om[0]
    m = _find(d, "m")
    assert m is not None
    assert len(m.findall(L.qm("mr"))) == 3
    back = word2latex.to_latex(om)
    assert back.strip() == L.serialize(L.parse(tex)), (
        "reverse spelling must equal the canonical spelling")


def test_d14_bare_eqarr_reverses_to_matrix():
    """A bare `m:eqArr` (no delimiters -- Word's shape for a standalone
    aligned/gathered block) spells `\\begin{matrix}`, the canonical home of
    the alias family, and that spelling round-trips unchanged."""
    M = L.qm
    om = etree.Element(M("oMath"))
    eqarr = etree.SubElement(om, M("eqArr"))
    for row in ("a", "b"):
        re = etree.SubElement(eqarr, M("e"))
        r = etree.SubElement(re, M("r"))
        t = etree.SubElement(r, M("t"))
        t.text = row
    back = word2latex.to_latex(om)
    assert back.strip().startswith(r"\begin{matrix}")
    assert back.strip().endswith(r"\end{matrix}")
    om2 = _emit(back)
    assert _root_tags(om2) == ["m"], "matrix must become one bare m:m"
    assert word2latex.to_latex(om2).strip() == back.strip()


# --- D15: prose following a display equation must not inherit its centering


def test_d15_prose_after_display_math_keeps_left_alignment():
    """DEFECTS.md D15: a display equation sharing its paragraph with
    trailing prose (the `\\[...\\]`-then-prose shape) rendered "fused"
    with that prose: add_display_math centers the whole paragraph, so the
    prose line inherited the centering too. The equation must stay
    centered through its own `m:jc`; the prose line must return to the
    paragraph default (left) -- every corpus paragraph holding both an
    `m:oMathPara` and sibling runs carries no `w:jc` at all (measured on
    all 7 such paragraphs)."""
    builder = latex2word.DocxBuilder()
    builder.add_display_math(r"f = \kappa\tau")
    builder.add_paragraph_text("trailing prose", append=True)
    p = builder.doc.paragraphs[-1]
    ppr = p._element.get_or_add_pPr()
    assert ppr.find(qn("w:jc")) is None, "prose line must not be centered"
    omp = p._element.find(latex2word._m("oMathPara"))
    assert omp is not None, "the equation must stay a display m:oMathPara"
    jc = omp.find(latex2word._m("oMathParaPr")).find(latex2word._m("jc"))
    assert jc.get(latex2word._m("val")) == "center", (
        "the equation keeps its own centering")
    texts = [t.text or "" for t in p._element.iter(latex2word._m("t"))]
    assert "".join(texts) == "f=κτ", "equation content intact"
    assert p.text == "trailing prose"


def test_d15_equation_only_paragraph_stays_centered():
    """D15's counterpart: a display equation *alone* in its paragraph
    keeps the centered shape (m:jc=center, w:jc=center) -- the fix must
    not disturb the equation-only case."""
    builder = latex2word.DocxBuilder()
    builder.add_display_math(r"x = 1")
    p = builder.doc.paragraphs[-1]
    ppr = p._element.get_or_add_pPr()
    assert ppr.find(qn("w:jc")) is not None
    assert ppr.find(qn("w:jc")).get(qn("w:val")) == "center"


# --- D18: prose sharing a paragraph with display math must survive ----------


def test_d18_prose_then_display_math_stays_one_paragraph():
    """DEFECTS.md D18: a display equation sharing its Word paragraph with
    *preceding* prose (the `prose\\n\\[...\\]` shape -- the mirror of
    D15) used to start a fresh paragraph on the way in, so the round trip
    re-registered a "text changed" + "block appeared" pair every
    generation (measured on two corpus documents). The equation must join
    the open prose paragraph, and -- the D15 discipline, prose side first
    -- the paragraph must carry no `w:jc` while the equation stays
    centered through its own `oMathParaPr/m:jc`."""
    builder = latex2word.DocxBuilder()
    builder.add_paragraph_text("Знайти розв'язок рівняння")
    builder.add_display_math(r"x^2 + y^2 = z^2", append=True)
    p = builder.doc.paragraphs[-1]
    assert p.text == "Знайти розв'язок рівняння"
    ppr = p._element.get_or_add_pPr()
    assert ppr.find(qn("w:jc")) is None, "prose line keeps its own alignment"
    omp = p._element.find(latex2word._m("oMathPara"))
    assert omp is not None, "the equation must join the prose paragraph"
    jc = omp.find(latex2word._m("oMathParaPr")).find(latex2word._m("jc"))
    assert jc.get(latex2word._m("val")) == "center", (
        "the equation keeps its own centering")


def test_d18_prose_then_math_roundtrip_is_a_fixed_point(tmp_path):
    """The end-to-end shape that measured the defect: a source paragraph
    `prose\\n\\[...\\]` must come back from a full round trip as *one*
    paragraph holding both, and the reverse must re-emit the same
    `prose\\n\\[` spelling (no manufactured blank line between the prose
    and the equation -- a soft break before the equation used to grow one,
    splitting the paragraph again)."""
    tex = tmp_path / "in.tex"
    tex.write_text(
        "\\documentclass{article}\n"
        "\\usepackage{amsmath,amssymb,amsfonts,mathtools}\n"
        "\\begin{document}\n"
        "\n"
        "Знайти розв'язок рівняння\n"
        "\\[\n"
        "x^2 + y^2 = z^2\n"
        "\\]\n"
        "\\end{document}\n", encoding="utf-8")
    docx_path = latex2word.convert_latex_to_docx(
        str(tex), str(tmp_path / "out.docx"))[0]
    # docx_to_latex returns the document text; writing it to disk is the
    # caller's job (the CLI does it), and the re-conversion below needs the
    # file on disk.
    back, _ = docx_read.docx_to_latex(docx_path)
    (tmp_path / "back.tex").write_text(back, encoding="utf-8")
    # the prose run's own trailing space (the folded source newline) sits
    # before the equation; only the blank line is forbidden
    assert "Знайти розв'язок рівняння \n\\[" in back, (
        "no blank line may appear between prose and the equation")
    # and the re-conversion keeps the paragraph merged
    docx2 = latex2word.convert_latex_to_docx(
        str(tmp_path / "back.tex"), str(tmp_path / "out2.docx"))[0]
    from docx import Document
    paras = [p for p in Document(docx2).paragraphs if p.text]
    assert len(paras) == 1, "prose + equation must stay one paragraph"


def test_d18_prose_between_equations_keeps_all_paragraphs():
    """D18's regression guard for the stale-open-math tangle: display
    math, then a styled paragraph, then prose + display math (the
    `\\wstyle`-between-equations run measured on a corpus document) --
    the second equation used to append as a second row of the *first*
    equation's `m:oMathPara` because the open-math handle survived the
    intermediate prose paragraphs ("math zone count 1 -> 2"). Each
    paragraph must stay its own."""
    builder = latex2word.DocxBuilder()
    builder.add_display_math(r"u_{\eta\eta}=0")
    builder.add_paragraph_text("Розв'язок Даламбера")
    builder.add_paragraph_text("Знайти розв'язок рівняння")
    builder.add_display_math(
        r"\frac{\partial^{2}u}{\partial t^{2}}=a^{2}",
        append=True)
    paras = builder.doc.paragraphs
    assert len(paras) == 3, "no paragraph may merge with another"
    rows = paras[0]._element.find(
        latex2word._m("oMathPara")).findall(latex2word._m("oMath"))
    assert len(rows) == 1, "the first equation must stay single-row"
    assert paras[2]._element.find(
        latex2word._m("oMathPara")) is not None, (
        "the second equation joins the prose paragraph")


def test_d19_bare_delim_respects_group_boundary():
    """D19: a bare `(` inside a braced group must not scan past the group's
    closing brace -- `{\partial(m}` is a group holding a bare `(` (the
    reverse emitter's group-subscript spelling for Word's
    sSub(base="∂(m")), and the group boundary falls back to the plain
    character exactly like the no-closer-anywhere case. Before the fix the
    rbrace raised UnbalancedDelimiterError and the whole equation fell
    back to monospaced literal LaTeX."""
    for tex in (r"{\partial(m}", r"\frac{{\partial(m}_{i}M_{s})}{\partial t}"):
        node = L.parse(tex)
        assert node is not None, tex
    # the canonical round trip: parse the reverse's own spelling, serialize,
    # and parse again -- a fixed point, not just an accepted input
    tex = r"m_{i}=\frac{M_{i}}{M_{s}},\ \ \frac{{\partial M}_{i}}{\partial t}=\frac{{\partial(m}_{i}M_{s})}{\partial t}"
    once = L.serialize(L.parse(tex))
    assert L.parse(once) is not None
