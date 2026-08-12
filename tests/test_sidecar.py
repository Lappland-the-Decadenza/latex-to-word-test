"""Synthetic characterization tests for detached Word object storage."""

import base64
import hashlib
import io
import json
import zipfile

import docx
import pytest

from latexword.docx.package import validate_docx_package
from latexword.document.identity import NodeId
from latexword.sidecar import ObjectStore
from latexword.sidecar.capture import content_types, relationships
from latexword.sidecar.paths import resolve_target, validate_relative


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
    "/A8AAwMBgFZ2q8sAAAAASUVORK5CYII="
)


def _picture_document(path):
    document = docx.Document()
    document.add_picture(io.BytesIO(_PNG))
    document.save(path)
    return document


def test_sidecar_paths_reject_escape_and_external_targets():
    assert resolve_target("word/document.xml", "media/image.png") == "word/media/image.png"
    assert resolve_target("word/charts/chart1.xml", "../embeddings/data.xlsx") == "word/embeddings/data.xlsx"
    assert resolve_target("word/document.xml", "../../outside.xml") is None
    assert resolve_target("word/document.xml", "/absolute.xml") is None
    assert resolve_target("word/document.xml", "https://example.test/x") is None

    assert validate_relative("parts/obj0001/part0000.png") == "parts/obj0001/part0000.png"
    for value in ("", "../outside", "/absolute", "parts\\..\\outside"):
        with pytest.raises(ValueError):
            validate_relative(value)


def test_capture_restores_relationship_closure_and_payload_bytes(tmp_path):
    source_path = tmp_path / "source.docx"
    source_document = _picture_document(source_path)
    drawing = source_document.paragraphs[0]._p.find(".//" + W + "drawing")
    assert drawing is not None

    tex_path = tmp_path / "shadow.tex"
    with zipfile.ZipFile(source_path) as archive:
        store = ObjectStore.for_write(tex_path, archive)
        object_id = store.capture(drawing, kind="drawing", context="inline")
        store.close()

    loaded = ObjectStore.for_read(tex_path)
    assert loaded is not None
    assert loaded.is_inline(object_id)
    manifest = json.loads((tmp_path / "shadow.objects" / "manifest.json").read_text())
    assert manifest["source_sha256"] == hashlib.sha256(source_path.read_bytes()).hexdigest()
    captured = manifest["objects"][0]
    assert captured["id"] == object_id
    assert captured["relationships"]
    assert captured["parts"]
    payload = tmp_path / "shadow.objects" / captured["parts"][0]["path"]
    assert payload.read_bytes() == _PNG

    assert ObjectStore.for_read(tex_path, source_path) is not None
    other_source = tmp_path / "other.docx"
    other_source.write_bytes(b"different source")
    with pytest.raises(ValueError, match="source hash mismatch"):
        ObjectStore.for_read(tex_path, other_source)

    restored = docx.Document()
    paragraph = restored.add_paragraph()
    assert loaded.restore(restored.part, object_id, paragraph=paragraph)
    loaded.close()
    output_path = tmp_path / "restored.docx"
    restored.save(output_path)
    assert validate_docx_package(output_path) == []


def test_package_part_capture_preserves_content_types_and_relationships(tmp_path):
    source_path = tmp_path / "source.docx"
    document = docx.Document()
    document.save(source_path)
    tex_path = tmp_path / "shadow.tex"

    with zipfile.ZipFile(source_path) as archive:
        get_type = content_types(archive)
        assert get_type("word/document.xml").endswith("document.main+xml")
        assert relationships(archive, "word/document.xml")
        store = ObjectStore.for_write(tex_path, archive)
        object_id = store.capture_package_part(
            "word/settings.xml",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings",
        )
        store.close()

    loaded = ObjectStore.for_read(tex_path)
    assert loaded is not None
    assert loaded._by_id[object_id].package_source == "word/settings.xml"
    loaded.close()


def test_typed_attachments_are_content_addressed_and_owner_bound(tmp_path):
    source_path = tmp_path / "source.docx"
    docx.Document().save(source_path)
    tex_path = tmp_path / "shadow.tex"
    with zipfile.ZipFile(source_path) as archive:
        store = ObjectStore.for_write(tex_path, archive)
        first = store.attach(
            "paragraph-style", {"style_id": "BodyText"},
            owner_id=NodeId.allocate(0), owner_semantic_hash="a" * 64,
            ordinal=0, content_type="application/json",
        )
        second = store.attach(
            "paragraph-style", {"style_id": "BodyText"},
            owner_id=NodeId.allocate(1), owner_semantic_hash="b" * 64,
            ordinal=1, content_type="application/json",
        )
        store.close()

    assert first.path == second.path
    loaded = ObjectStore.for_read(tex_path)
    assert loaded is not None
    assert len(loaded.attachments_for(NodeId.allocate(0))) == 1
    assert loaded.attachments_at(1)[0].owner_id == NodeId.allocate(1)
    assert json.loads(loaded.attachment_payload(first)) == {
        "style_id": "BodyText"
    }
    loaded.close()


def test_sidecar_rejects_ambiguous_or_incomplete_relationship_graph(tmp_path):
    source_path = tmp_path / "source.docx"
    source_document = _picture_document(source_path)
    drawing = source_document.paragraphs[0]._p.find(".//" + W + "drawing")
    tex_path = tmp_path / "shadow.tex"
    with zipfile.ZipFile(source_path) as archive:
        store = ObjectStore.for_write(tex_path, archive)
        store.capture(drawing, kind="drawing", context="inline")
        store.close()

    manifest_path = tmp_path / "shadow.objects" / "manifest.json"
    manifest = json.loads(manifest_path.read_text())
    relationships = manifest["objects"][0]["relationships"]
    relationships.append(dict(relationships[0]))
    manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
    with pytest.raises(ValueError, match="ambiguous"):
        ObjectStore.for_read(tex_path)

    manifest = json.loads(manifest_path.read_text())
    relationships = manifest["objects"][0]["relationships"]
    relationships.pop()
    relationships[0]["part"] = "missing.xml"
    manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
    with pytest.raises(ValueError, match="incomplete"):
        ObjectStore.for_read(tex_path)
